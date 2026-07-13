import os
import logging
import img2pdf
import zipfile
import subprocess
from pdf2docx import Converter
from docx import Document
from PIL import Image
import fitz
import pytesseract

logger = logging.getLogger(__name__)

MAX_WATERMARK_LEN = 200

def create_pdf_from_images(image_paths: list, output_path: str):
    """Rasmlarni bitta PDF faylga birlashtirish"""
    if not image_paths:
        return False

    # Rasmlarni PDF uchun to'g'rilash (masalan formati xato bo'lsa)
    valid_images = []
    for img_path in image_paths:
        try:
            with Image.open(img_path) as img:
                img.verify() # formatni tekshirish
            valid_images.append(img_path)
        except Exception:
            pass

    if not valid_images:
        return False

    with open(output_path, "wb") as f:
        f.write(img2pdf.convert(valid_images))
    return True

def create_zip_from_files(file_paths: list, output_path: str):
    """Fayllarni ZIP arxivga joylash"""
    if not file_paths:
        return False

    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file in file_paths:
            # Fayl nomidagi '___' dan keyingi asl nomni ajratib olish (yoki shunchaki basename)
            basename = os.path.basename(file)
            arcname = basename.split("___")[-1] if "___" in basename else basename
            zipf.write(file, arcname)
    return True

def text_to_docx(text: str, output_path: str):
    """Matnni DOCX ga aylantirish"""
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)
    return True

def pdf_to_docx(pdf_path: str, output_path: str):
    """PDF faylni DOCX ga o'girish (Real Document Converter xususiyati)"""
    try:
        cv = Converter(pdf_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        return True
    except Exception as e:
        logger.exception("PDF to DOCX Error: %s", e)
        return False

def merge_pdfs(pdf_paths: list, output_path: str):
    """Bir nechta PDF fayllarni birlashtirish"""
    try:
        result = fitz.open()
        for pdf in pdf_paths:
            with fitz.open(pdf) as mfile:
                result.insert_pdf(mfile)
        result.save(output_path)
        result.close()
        return True
    except Exception as e:
        logger.exception("Merge PDF Error: %s", e)
        return False

def split_pdf(pdf_path: str, output_path: str, start_page: int, end_page: int):
    """PDF ni qirqish (start_page dan end_page gacha, 1-indexed, inclusive)"""
    doc = None
    result = None
    try:
        doc = fitz.open(pdf_path)
        page_count = doc.page_count
        if start_page < 1 or end_page < start_page or end_page > page_count:
            logger.warning(
                "Split PDF range out of bounds: start=%s end=%s page_count=%s",
                start_page, end_page, page_count,
            )
            return False
        result = fitz.open()
        # Pages are 0-indexed in PyMuPDF
        result.insert_pdf(doc, from_page=start_page - 1, to_page=end_page - 1)
        result.save(output_path)
        return True
    except Exception as e:
        logger.exception("Split PDF Error: %s", e)
        return False
    finally:
        if result is not None:
            result.close()
        if doc is not None:
            doc.close()

def add_watermark(pdf_path: str, output_path: str, watermark_text: str):
    """PDF ning butun yuzasi bo'ylab takrorlanadigan diagonal (45°) yarim-shaffof watermark.

    Matn butun sahifa bo'ylab to'r (grid) shaklida joylashtiriladi va markaz atrofida
    45° ga buriladi. fitz.Page.insert_text() faqat 0/90/180/270 ni qabul qilgani uchun
    ixtiyoriy burchak TextWriter + morph (Matrix) orqali beriladi.
    """
    watermark_text = (watermark_text or "").strip()[:MAX_WATERMARK_LEN]
    if not watermark_text:
        return False
    doc = None
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            rect = page.rect
            width, height = rect.width, rect.height
            fontsize = max(18, int(width / 22))
            text_len = fitz.get_text_length(watermark_text, fontsize=fontsize)
            step_x = text_len + fontsize * 2.5
            step_y = fontsize * 4.5
            tw = fitz.TextWriter(rect, color=(0.5, 0.5, 0.5))
            # 45° burilishdan keyin burchaklar ham to'lishi uchun sahifadan tashqariga chiqib qo'yamiz
            y = -height
            while y < 2 * height:
                x = -width
                while x < 2 * width:
                    try:
                        tw.append(fitz.Point(x, y), watermark_text, fontsize=fontsize)
                    except Exception:
                        pass
                    x += step_x
                y += step_y
            tw.write_text(page, opacity=0.18, morph=(fitz.Point(width / 2, height / 2), fitz.Matrix(45)))
        doc.save(output_path)
        return True
    except Exception as e:
        logger.exception("Watermark Error: %s", e)
        return False
    finally:
        if doc is not None:
            doc.close()

def office_to_pdf(input_path: str, out_dir: str, timeout: int = 120):
    """Convert docx, xlsx, pptx to pdf using LibreOffice headless"""
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", input_path, "--outdir", out_dir],
            check=True,
            timeout=timeout,
        )
        base = os.path.basename(input_path)
        name, _ = os.path.splitext(base)
        pdf_path = os.path.join(out_dir, f"{name}.pdf")
        return pdf_path if os.path.exists(pdf_path) else None
    except subprocess.TimeoutExpired as e:
        logger.error("Office to PDF timed out after %ss: %s", timeout, e)
        return None
    except Exception as e:
        logger.exception("Office to PDF Error: %s", e)
        return None

def image_to_text_docx(image_path: str, output_path: str, lang: str = 'uzb+rus+eng'):
    """Extract text from image using Tesseract and save to DOCX"""
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang=lang)
        doc = Document()
        doc.add_paragraph(text)
        doc.save(output_path)
        return True
    except Exception as e:
        logger.exception("OCR Error: %s", e)
        return False
