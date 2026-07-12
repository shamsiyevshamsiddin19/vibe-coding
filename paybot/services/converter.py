import os
import img2pdf
import zipfile
import subprocess
from pdf2docx import Converter
from docx import Document
from PIL import Image
import fitz
import pytesseract

def create_pdf_from_images(image_paths: list, output_path: str):
    """Rasmlarni bitta PDF faylga birlashtirish"""
    if not image_paths:
        return False
    
    # Rasmlarni PDF uchun to'g'rilash (masalan formati xato bo'lsa)
    valid_images = []
    for img_path in image_paths:
        try:
            with Image.open(img_path) as img:
                img.verify() # formatni tekshirish
            valid_images.append(img_path)
        except Exception:
            pass
            
    if not valid_images:
        return False

    with open(output_path, "wb") as f:
        f.write(img2pdf.convert(valid_images))
    return True

def create_zip_from_files(file_paths: list, output_path: str):
    """Fayllarni ZIP arxivga joylash"""
    if not file_paths:
        return False
        
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file in file_paths:
            # Fayl nomidagi '___' dan keyingi asl nomni ajratib olish (yoki shunchaki basename)
            basename = os.path.basename(file)
            arcname = basename.split("___")[-1] if "___" in basename else basename
            zipf.write(file, arcname)
    return True

def text_to_docx(text: str, output_path: str):
    """Matnni DOCX ga aylantirish"""
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)
    return True

def pdf_to_docx(pdf_path: str, output_path: str):
    """PDF faylni DOCX ga o'girish (Real Document Converter xususiyati)"""
    try:
        cv = Converter(pdf_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        return True
    except Exception as e:
        print(f"PDF to DOCX Error: {e}")
        return False

import fitz

def merge_pdfs(pdf_paths: list, output_path: str):
    """Bir nechta PDF fayllarni birlashtirish"""
    try:
        result = fitz.open()
        for pdf in pdf_paths:
            with fitz.open(pdf) as mfile:
                result.insert_pdf(mfile)
        result.save(output_path)
        result.close()
        return True
    except Exception as e:
        print(f"Merge PDF Error: {e}")
        return False

def split_pdf(pdf_path: str, output_path: str, start_page: int, end_page: int):
    """PDF ni qirqish (start_page dan end_page gacha)"""
    try:
        doc = fitz.open(pdf_path)
        result = fitz.open()
        # Pages are 0-indexed in PyMuPDF
        result.insert_pdf(doc, from_page=start_page-1, to_page=end_page-1)
        result.save(output_path)
        doc.close()
        result.close()
        return True
    except Exception as e:
        print(f"Split PDF Error: {e}")
        return False

def add_watermark(pdf_path: str, output_path: str, watermark_text: str):
    """PDF ga watermark qo'shish"""
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            rect = page.rect
            p = fitz.Point(rect.width / 4, rect.height / 2)
            page.insert_text(p, watermark_text, fontsize=50, color=(0.5, 0.5, 0.5), fill_opacity=0.3, rotate=45)
        doc.save(output_path)
        doc.close()
        return True
    except Exception as e:
        print(f"Watermark Error: {e}")
        return False

def office_to_pdf(input_path: str, out_dir: str):
    """Convert docx, xlsx, pptx to pdf using LibreOffice headless"""
    try:
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", input_path, "--outdir", out_dir], check=True)
        base = os.path.basename(input_path)
        name, _ = os.path.splitext(base)
        pdf_path = os.path.join(out_dir, f"{name}.pdf")
        return pdf_path if os.path.exists(pdf_path) else None
    except Exception as e:
        print(f"Office to PDF Error: {e}")
        return None

def image_to_text_docx(image_path: str, output_path: str, lang: str = 'uzb+rus+eng'):
    """Extract text from image using Tesseract and save to DOCX"""
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang=lang)
        doc = Document()
        doc.add_paragraph(text)
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"OCR Error: {e}")
        return False

