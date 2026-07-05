"""Markdown matnidan akademik DOCX fayl yaratadi.

Qo'llab-quvvatlanadi: sarlavhalar, ro'yxatlar, ichki **qalin** matn,
jadvallar (Word Table, ramkali, sarlavha qatori bo'yalgan), formulalar
($$...$$ — markazda, Cambria Math), grafiklar (```chart JSON``` → PNG).
"""
from __future__ import annotations
import logging
import os
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docs.md_blocks import parse_blocks, split_bold_runs, strip_md, clean_ctrl

logger = logging.getLogger(__name__)

_FONT = "Times New Roman"


def _set_page_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def _add_page_numbers(doc: Document):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run._r.append(fld)
        ins = OxmlElement("w:instrText")
        ins.text = " PAGE "
        run._r.append(ins)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        run._r.append(fld2)


def _para(doc: Document, text: str, bold: bool = False, size: int = 14,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before: float = 0,
          space_after: float = 6, italic: bool = False, font: str = _FONT):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(21)  # 1.5 interval
    for seg, seg_bold in split_bold_runs(text):
        run = p.add_run(clean_ctrl(seg))
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold or seg_bold
        run.font.italic = italic
    return p


def _heading(doc: Document, text: str, level: int = 1):
    sizes = {1: 16, 2: 15, 3: 14}
    _para(doc, text.upper() if level == 1 else text,
          bold=True, size=sizes.get(level, 14),
          align=WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT,
          space_before=12, space_after=8)


def _list_item(doc: Document, text: str, style: str):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for seg, seg_bold in split_bold_runs(text):
        run = p.add_run(clean_ctrl(seg))
        run.font.name = _FONT
        run.font.size = Pt(14)
        run.font.bold = seg_bold


def _shade_cell(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if ri == 0
                           else WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(strip_md(cell_text))
            run.font.name = _FONT
            run.font.size = Pt(12)
            run.font.bold = ri == 0
            if ri == 0:
                _shade_cell(cell, "DEEAF6")  # och ko'k sarlavha qatori
    # Jadvaldan keyin kichik bo'shliq
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


_OMML_NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'


def _add_formula(doc: Document, formula: str):
    """Formulani Word'ning haqiqiy (OMML) tenglamasi sifatida joylaydi.
    LaTeX → MathML → OMML; xatoda Unicode-kursiv fallback."""
    try:
        import latex2mathml.converter
        import mathml2omml
        from docx.oxml import parse_xml
        mathml = latex2mathml.converter.convert(formula)
        omml = mathml2omml.convert(mathml)
        # Namespace deklaratsiyasini qo'shamiz (parse_xml talab qiladi)
        omml = omml.replace("<m:oMath>", f"<m:oMath {_OMML_NS}>", 1)
        element = parse_xml(omml)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(8)
        p._p.append(element)
        return
    except Exception as e:
        logger.debug("LaTeX→OMML muvaffaqiyatsiz (%s) — Unicode fallback", e)
    _para(doc, formula, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=6, space_after=8, italic=True, font="Cambria Math")


_MINISTRY = {
    "uz": "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI",
    "ru": "МИНИСТЕРСТВО ВЫСШЕГО ОБРАЗОВАНИЯ, НАУКИ И ИННОВАЦИЙ РЕСПУБЛИКИ УЗБЕКИСТАН",
    "en": "MINISTRY OF HIGHER EDUCATION, SCIENCE AND INNOVATIONS OF THE REPUBLIC OF UZBEKISTAN",
}
_CITY = {"uz": "Toshkent", "ru": "Ташкент", "en": "Tashkent"}


def add_title_page(doc: Document, meta: dict):
    """Standart akademik titul varag'i (kod bilan — aniq joylashuv)."""
    import datetime
    lang = meta.get("language", "uz")
    center = WD_ALIGN_PARAGRAPH.CENTER

    def line(text, size=14, bold=False, before=0, after=6):
        if text:
            _para(doc, text, bold=bold, size=size, align=center,
                  space_before=before, space_after=after)

    def gap(n=1):
        for _ in range(n):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)

    line(_MINISTRY.get(lang, _MINISTRY["uz"]), size=13, bold=True)
    otm = (meta.get("otm") or "").strip()
    if otm:
        line(otm.upper(), size=14, bold=True, before=6)
    gap(3)
    fan = (meta.get("fan") or "").strip()
    if fan and fan != "—":
        line(f"«{fan}» fanidan" if lang == "uz" else fan, size=14)
    gap(1)
    line(meta.get("doc_label", "").upper(), size=22, bold=True, before=10, after=10)
    topic = (meta.get("topic") or "").strip()
    line(f"Mavzu: {topic}" if lang == "uz" else f"Тема: {topic}" if lang == "ru"
         else f"Topic: {topic}", size=16, bold=True, after=10)
    gap(4)

    # O'ng tomonda bajaruvchi ma'lumotlari
    right = WD_ALIGN_PARAGRAPH.RIGHT
    author = (meta.get("author") or "").strip()
    group = (meta.get("group") or "").strip()
    supervisor = (meta.get("supervisor") or "").strip()
    if author and author != "—":
        who = f"Bajardi: {author}" + (f" ({group})" if group and group != "—" else "")
        _para(doc, who, size=14, align=right, space_after=4)
    if supervisor and supervisor != "—":
        _para(doc, f"Ilmiy rahbar: {supervisor}", size=14, align=right, space_after=4)
    gap(6)
    line(f"{_CITY.get(lang, 'Toshkent')} — {datetime.date.today().year}",
         size=14, bold=True)
    doc.add_page_break()


def _add_chart(doc: Document, spec: dict, temp_files: list[str]) -> bool:
    try:
        from docs.charts import render_chart
        fd, png = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        render_chart(spec, png)
        temp_files.append(png)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(png, width=Cm(14.5))
        return True
    except Exception as e:
        logger.warning("Grafik chizib bo'lmadi (%s) — o'tkazib yuborildi", e)
        title = str(spec.get("title", "")).strip() if isinstance(spec, dict) else ""
        if title:
            _para(doc, title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        return False


def build_docx(markdown_text: str, output_path: str, title: str = "",
               meta: dict | None = None) -> str:
    doc = Document()
    _set_page_margins(doc)
    _add_page_numbers(doc)

    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(14)

    if meta:
        try:
            # Titulda sahifa raqami ko'rinmasin
            doc.sections[0].different_first_page_header_footer = True
        except Exception:
            pass
        add_title_page(doc, meta)

    temp_files: list[str] = []
    for block in parse_blocks(markdown_text):
        kind = block[0]
        if kind == "h1":
            _heading(doc, block[1], level=1)
        elif kind == "h2":
            _heading(doc, block[1], level=2)
        elif kind == "h3":
            _heading(doc, block[1], level=3)
        elif kind == "bullet":
            _list_item(doc, block[1], "List Bullet")
        elif kind == "number":
            _list_item(doc, block[1], "List Number")
        elif kind == "bold":
            _para(doc, block[1], bold=True)
        elif kind == "table":
            _add_table(doc, block[1])
        elif kind == "formula":
            _add_formula(doc, block[1])
        elif kind == "chart":
            _add_chart(doc, block[1], temp_files)
        else:
            _para(doc, block[1])

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    for f in temp_files:
        try:
            os.remove(f)
        except OSError:
            pass
    return output_path
