"""Krossvord generatsiya: JSON so'zlar → grid + chiroyli rasm + to'liq DOCX.

- Joylashuv qoidalari to'g'ri (qo'shni kataklar tegmaydi, faqat kesishish)
- Har safar boshqa shakl: so'z tartibi va kesishish tanlovi tasodifiy
- 5 xil rang-mavzu — har buyurtmada tasodifiy tanlanadi
- DOCX ichida: savol kataklari rasmi + tariflar + alohida sahifada javoblar
"""
from __future__ import annotations
import json
import os
import random
import re
import tempfile
from PIL import Image, ImageDraw, ImageFont

from docs.json_utils import parse_json_objects

CELL = 56          # piksel, katakcha o'lchami
GRID = 23          # maksimal grid hajmi
MARGIN = 36
FONT_SIZE = 26

# Rang-mavzular: katak, chegara, harf, raqam
THEMES = [
    {"name": "ko'k", "bg": (255, 255, 255), "cell": (239, 246, 255),
     "border": (30, 64, 175), "text": (30, 58, 138), "num": (220, 38, 38)},
    {"name": "yashil", "bg": (255, 255, 255), "cell": (236, 253, 245),
     "border": (6, 95, 70), "text": (6, 78, 59), "num": (190, 18, 60)},
    {"name": "binafsha", "bg": (255, 255, 255), "cell": (245, 243, 255),
     "border": (109, 40, 217), "text": (76, 29, 149), "num": (217, 70, 39)},
    {"name": "terrakota", "bg": (255, 255, 255), "cell": (255, 247, 237),
     "border": (154, 52, 18), "text": (124, 45, 18), "num": (21, 94, 117)},
    {"name": "feruza", "bg": (255, 255, 255), "cell": (240, 253, 250),
     "border": (15, 118, 110), "text": (17, 94, 89), "num": (190, 24, 93)},
]

_FONT_PATHS = [
    r"C:\Windows\Fonts\arialbd.ttf",
    "/usr/share/fonts/dejavu-sans-fonts/DejaVuSans-Bold.ttf",
    "/usr/share/fonts/liberation-sans/LiberationSans-Bold.ttf",
]


def _font(size: int):
    for p in _FONT_PATHS:
        if os.path.isfile(p):
            try:
                return ImageFont.truetype(p, size)
            except Exception:
                continue
    return ImageFont.load_default()


def _clean_word(word: str) -> str:
    """Katta harf, apostrof va harf bo'lmaganlarni olib tashlash."""
    w = str(word).upper()
    w = w.replace("ʼ", "").replace("'", "").replace("’", "").replace("`", "")
    return "".join(ch for ch in w if ch.isalpha())


def _parse_words(json_text: str) -> list[dict]:
    out, seen = [], set()
    for d in parse_json_objects(json_text):
        if "word" not in d or "clue" not in d:
            continue
        w = _clean_word(d["word"])
        if 3 <= len(w) <= 14 and w not in seen:
            seen.add(w)
            out.append({"word": w, "clue": str(d["clue"]).strip()})
    return out


class Crossword:
    def __init__(self, words: list[dict]):
        self.words = words
        self.grid: dict[tuple, str] = {}
        self.placed: list[dict] = []

    # ─── Joylashuv qoidalari ───

    def _fits(self, word: str, row: int, col: int, horiz: bool) -> bool:
        dr, dc = (0, 1) if horiz else (1, 0)
        if row < 0 or col < 0:
            return False
        if horiz and col + len(word) > GRID:
            return False
        if not horiz and row + len(word) > GRID:
            return False
        # So'zdan oldingi va keyingi katak bo'sh bo'lishi shart
        if (row - dr, col - dc) in self.grid:
            return False
        if (row + dr * len(word), col + dc * len(word)) in self.grid:
            return False
        crossings = 0
        for i, ch in enumerate(word):
            r, c = row + dr * i, col + dc * i
            cur = self.grid.get((r, c))
            if cur is not None:
                if cur != ch:
                    return False
                crossings += 1
                continue
            # Yangi katak: perpendikulyar qo'shnilar bo'sh bo'lsin
            if horiz:
                if (r - 1, c) in self.grid or (r + 1, c) in self.grid:
                    return False
            else:
                if (r, c - 1) in self.grid or (r, c + 1) in self.grid:
                    return False
        return crossings >= 1  # kamida bitta kesishish (birinchi so'zdan tashqari)

    def _place(self, entry: dict, row: int, col: int, horiz: bool):
        word = entry["word"]
        dr, dc = (0, 1) if horiz else (1, 0)
        for i, ch in enumerate(word):
            self.grid[(row + dr * i, col + dc * i)] = ch
        self.placed.append({"word": word, "clue": entry["clue"],
                            "row": row, "col": col, "horiz": horiz})

    def build(self):
        if self.placed or not self.words:
            return
        # Uzunlik bo'yicha, lekin tasodifiy aralashuv bilan — har safar boshqa shakl
        words = sorted(self.words, key=lambda w: (-len(w["word"]), random.random()))
        first = words[0]
        horiz0 = random.random() < 0.5
        w0 = first["word"]
        if horiz0:
            r0, c0 = GRID // 2, max(0, (GRID - len(w0)) // 2)
        else:
            r0, c0 = max(0, (GRID - len(w0)) // 2), GRID // 2
        dr, dc = (0, 1) if horiz0 else (1, 0)
        for i, ch in enumerate(w0):
            self.grid[(r0 + dr * i, c0 + dc * i)] = ch
        self.placed.append({"word": w0, "clue": first["clue"],
                            "row": r0, "col": c0, "horiz": horiz0})

        # Ikki o'tishda joylashtiramiz (birinchi o'tishda sig'maganlar keyin urinadi)
        pending = words[1:]
        for _ in range(2):
            leftover = []
            for entry in pending:
                if not self._try_place(entry):
                    leftover.append(entry)
            pending = leftover
            if not pending:
                break
        # Sig'magan so'zlar krossvordga kirmaydi (tariflar ham chiqarilmaydi)

    def _try_place(self, entry: dict) -> bool:
        word = entry["word"]
        candidates = []
        for p in self.placed:
            pw = p["word"]
            pdr, pdc = (0, 1) if p["horiz"] else (1, 0)
            for ci, ch in enumerate(word):
                for pi, pc in enumerate(pw):
                    if ch != pc:
                        continue
                    cross_r = p["row"] + pdr * pi
                    cross_c = p["col"] + pdc * pi
                    horiz = not p["horiz"]
                    if horiz:
                        row, col = cross_r, cross_c - ci
                    else:
                        row, col = cross_r - ci, cross_c
                    if self._fits(word, row, col, horiz):
                        candidates.append((row, col, horiz))
        if not candidates:
            return False
        row, col, horiz = random.choice(candidates)
        self._place(entry, row, col, horiz)
        return True

    # ─── Raqamlash (standart krossvord uslubida) ───

    def _numbering(self) -> dict[tuple, int]:
        starts = sorted({(p["row"], p["col"]) for p in self.placed})
        nums = {cell: i + 1 for i, cell in enumerate(starts)}
        for p in self.placed:
            p["num"] = nums[(p["row"], p["col"])]
        return nums

    def _bounds(self):
        rows = [r for r, c in self.grid]
        cols = [c for r, c in self.grid]
        return min(rows), min(cols), max(rows), max(cols)

    # ─── Rasm ───

    def render_image(self, output_path: str, theme: dict,
                     show_answers: bool = False) -> str:
        self.build()
        if not self.grid:
            raise ValueError("Krossvord kataklari yaratilmadi")
        nums = self._numbering()

        r0, c0, r1, c1 = self._bounds()
        rows, cols = r1 - r0 + 1, c1 - c0 + 1
        width = cols * CELL + 2 * MARGIN
        height = rows * CELL + 2 * MARGIN

        img = Image.new("RGB", (width, height), theme["bg"])
        draw = ImageDraw.Draw(img)
        font = _font(FONT_SIZE)
        numfont = _font(15)

        for (r, c), ch in self.grid.items():
            x = (c - c0) * CELL + MARGIN
            y = (r - r0) * CELL + MARGIN
            draw.rounded_rectangle(
                [x + 1, y + 1, x + CELL - 2, y + CELL - 2],
                radius=9, fill=theme["cell"], outline=theme["border"], width=3)
            if show_answers:
                tw = draw.textlength(ch, font=font)
                draw.text((x + (CELL - tw) / 2, y + (CELL - FONT_SIZE) / 2 - 3),
                          ch, fill=theme["text"], font=font)

        for (r, c), num in nums.items():
            x = (c - c0) * CELL + MARGIN + 5
            y = (r - r0) * CELL + MARGIN + 2
            draw.text((x, y), str(num), fill=theme["num"], font=numfont)

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        img.save(output_path, "PNG", dpi=(150, 150))
        return output_path

    def clues(self) -> tuple[list, list]:
        horiz = sorted([(p["num"], p["word"], p["clue"])
                        for p in self.placed if p["horiz"]])
        vert = sorted([(p["num"], p["word"], p["clue"])
                       for p in self.placed if not p["horiz"]])
        return horiz, vert


# ─────────────────────── DOCX yasash ───────────────────────

def _build_docx(cw: Crossword, docx_path: str, question_png: str,
                answer_png: str, topic: str, theme: dict):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor as DocxColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    accent = DocxColor(*theme["border"])

    def para(text, size=13, bold=False, center=False, color=None, after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(after)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color
        return p

    def picture(path, width_cm=16.5):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Rasm juda baland bo'lsa kichraytiramiz (A4 ga sig'sin)
        from PIL import Image as PILImage
        with PILImage.open(path) as im:
            w, h = im.size
        width = min(width_cm, 24 * w / h)  # balandlik ~24 sm dan oshmasin
        p.add_run().add_picture(path, width=Cm(width))

    # 1-sahifa: sarlavha + savol kataklari + tariflar
    para("KROSSVORD", size=22, bold=True, center=True, color=accent, after=2)
    if topic:
        para(f"Mavzu: {topic}", size=14, bold=True, center=True, after=10)
    picture(question_png)

    horiz, vert = cw.clues()
    para("Gorizontal (→)", size=14, bold=True, color=accent, after=4)
    for num, word, clue in horiz:
        para(f"{num}. {clue} ({len(word)} harf)", size=12, after=2)
    para("", after=2)
    para("Vertikal (↓)", size=14, bold=True, color=accent, after=4)
    for num, word, clue in vert:
        para(f"{num}. {clue} ({len(word)} harf)", size=12, after=2)

    # 2-sahifa: javoblar
    doc.add_page_break()
    para("JAVOBLAR", size=20, bold=True, center=True, color=accent, after=8)
    picture(answer_png)
    para("Gorizontal: " + ", ".join(f"{n} — {w}" for n, w, _ in horiz),
         size=12, after=4)
    para("Vertikal: " + ", ".join(f"{n} — {w}" for n, w, _ in vert),
         size=12, after=4)

    os.makedirs(os.path.dirname(docx_path) or ".", exist_ok=True)
    doc.save(docx_path)


def build_crossword(json_text: str, png_path: str, docx_path: str,
                    topic: str = "") -> tuple[str, str]:
    words = _parse_words(json_text)
    if len(words) < 8:
        raise ValueError(f"Krossvord uchun so'zlar yetarli emas ({len(words)} ta)")

    theme = random.choice(THEMES)
    cw = Crossword(words)
    cw.build()
    if len(cw.placed) < max(8, len(words) // 3):
        # Juda kam so'z joylashdi — boshqa tasodifiy tartib bilan yana urinamiz
        best = cw
        for _ in range(4):
            alt = Crossword(words)
            alt.build()
            if len(alt.placed) > len(best.placed):
                best = alt
        cw = best
    if len(cw.placed) < 8:
        raise ValueError(f"Krossvord tuzilmadi (faqat {len(cw.placed)} so'z joylashdi)")

    cw.render_image(png_path, theme, show_answers=False)
    fd, answer_png = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    try:
        cw.render_image(answer_png, theme, show_answers=True)
        _build_docx(cw, docx_path, png_path, answer_png, topic, theme)
    finally:
        try:
            os.remove(answer_png)
        except OSError:
            pass
    return png_path, docx_path
