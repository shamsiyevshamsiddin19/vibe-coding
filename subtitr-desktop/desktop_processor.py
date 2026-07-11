#!/usr/bin/env python3
"""Local desktop subtitle processor for the Flutter app.

The script is intentionally standalone: it does not import the Telegram bot,
database, Oracle, Celery, or web pieces. It reads videos from the shared
workspace, creates results in the output folder, and prints JSON lines so
Flutter can show live progress.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import threading
import io
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

if sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", line_buffering=True)
if sys.stderr.encoding.lower() != "utf-8":
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", line_buffering=True)

try:
    from dotenv import load_dotenv
except Exception:  # pragma: no cover - optional dependency
    load_dotenv = None

KINO_DIR = Path.home() / "Videos"
if getattr(sys, 'frozen', False):
    # Installed app: the exe may live in a read-only location (Program Files),
    # so results go to a visible folder in Videos and working/support files go
    # to a per-user writable data directory.
    ROOT = Path(sys.executable).resolve().parent
    _DATA = Path(os.getenv("LOCALAPPDATA") or Path.home()) / "SubtitrDesktop"
    OUT_DIR = KINO_DIR / "Subtitr natijalar"
    TMP_DIR = _DATA / "Ishchi fayllar"
    MODEL_DIR = _DATA / "AI modellar"
    DICT_DIR = _DATA / "Lugatlar"
    CACHE_DIR = _DATA / "cache"
else:
    ROOT = Path(__file__).resolve().parent
    OUT_DIR = ROOT / "Tayyor natijalar"
    TMP_DIR = ROOT / "Ishchi fayllar"
    MODEL_DIR = ROOT / "AI modellar"
    DICT_DIR = ROOT / "Lugatlar"
    CACHE_DIR = ROOT / "cache"
ENV_FILE = ROOT / ".env"


def resolve_tool(name: str) -> str:
    """Prefer a tool bundled next to the executable, then a dev `tools/` folder,
    else fall back to PATH."""
    fname = name + (".exe" if os.name == "nt" else "")
    for base in (ROOT, ROOT / "tools"):
        exe = base / fname
        if exe.exists():
            return str(exe)
    found = shutil.which(name)
    return found or name


FFMPEG = resolve_tool("ffmpeg")
FFPROBE = resolve_tool("ffprobe")
YTDLP = resolve_tool("yt-dlp")

VIDEO_EXTS = {".mp4", ".mkv", ".mov", ".avi", ".webm", ".m4v"}
SUB_EXTS = {".srt", ".vtt"}

TARGET_LANG_NAMES = {
    "uz": "Uzbek",
    "en": "English",
    "ru": "Russian",
    "tr": "Turkish",
    "kk": "Kazakh",
    "tg": "Tajik",
    "ky": "Kyrgyz",
}

SOURCE_LANG_NAMES = {
    "auto": "auto-detected source language",
    "uz": "Uzbek",
    "en": "English",
    "ru": "Russian",
    "tr": "Turkish",
    "kk": "Kazakh",
    "tg": "Tajik",
    "ky": "Kyrgyz",
}

HELPERS = {
    "en": {
        "article": {"a", "an", "the"},
        "preposition": {
            "in", "on", "at", "to", "for", "from", "with", "by", "of",
            "about", "into", "onto", "over", "under", "after", "before",
            "between", "through", "during", "without", "against", "around",
            "off", "up", "down", "as",
        },
        "conjunction": {
            "and", "but", "or", "so", "because", "if", "when", "while",
            "that", "than", "though", "although", "since", "unless", "yet",
        },
    },
    "uz": {
        "komakchi": {
            "bilan", "uchun", "kabi", "singari", "qadar", "tomon",
            "orqali", "haqida", "keyin", "oldin", "song", "qarshi",
        },
        "boglovchi": {
            "va", "ham", "ammo", "lekin", "biroq", "yoki", "chunki",
            "agar", "yani", "balki",
        },
    },
    "ru": {
        "preposition": {
            "v", "na", "s", "k", "u", "o", "po", "za", "iz", "ot",
            "do", "dlya", "pod", "nad", "pri", "pro", "bez", "cherez",
        },
        "conjunction": {"i", "a", "no", "ili", "chto", "esli", "kogda", "kak"},
    },
}

COMMON_WORDS = {
    "hello": "salom",
    "hi": "salom",
    "yes": "ha",
    "no": "yo'q",
    "not": "emas",
    "i": "men",
    "you": "siz",
    "he": "u",
    "she": "u",
    "we": "biz",
    "they": "ular",
    "it": "u",
    "me": "menga",
    "my": "mening",
    "your": "sizning",
    "our": "bizning",
    "this": "bu",
    "that": "o'sha",
    "these": "bular",
    "those": "o'shalar",
    "is": "bo'ladi",
    "are": "bo'ladi",
    "was": "edi",
    "were": "edi",
    "be": "bo'lmoq",
    "have": "ega bo'lmoq",
    "has": "ega",
    "do": "qilmoq",
    "does": "qiladi",
    "did": "qildi",
    "will": "bo'ladi",
    "can": "qila oladi",
    "could": "qila olardi",
    "should": "kerak",
    "must": "shart",
    "go": "bormoq",
    "come": "kelmoq",
    "see": "ko'rmoq",
    "look": "qaramoq",
    "know": "bilmoq",
    "think": "o'ylamoq",
    "want": "xohlamoq",
    "need": "kerak bo'lmoq",
    "make": "qilmoq",
    "take": "olmoq",
    "give": "bermoq",
    "get": "olmoq",
    "say": "aytmoq",
    "tell": "aytmoq",
    "speak": "gapirmoq",
    "talk": "suhbatlashmoq",
    "work": "ishlamoq",
    "time": "vaqt",
    "day": "kun",
    "night": "tun",
    "man": "erkak",
    "woman": "ayol",
    "people": "odamlar",
    "friend": "do'st",
    "home": "uy",
    "house": "uy",
    "world": "dunyo",
    "life": "hayot",
    "love": "sevgi",
    "money": "pul",
    "good": "yaxshi",
    "bad": "yomon",
    "big": "katta",
    "small": "kichik",
    "new": "yangi",
    "old": "eski",
    "great": "zo'r",
    "right": "to'g'ri",
    "wrong": "noto'g'ri",
    "now": "hozir",
    "then": "keyin",
    "here": "bu yerda",
    "there": "u yerda",
    "why": "nega",
    "what": "nima",
    "who": "kim",
    "where": "qayerda",
    "when": "qachon",
    "how": "qanday",
    "and": "va",
    "but": "lekin",
    "or": "yoki",
    "because": "chunki",
    "for": "uchun",
    "with": "bilan",
    "in": "ichida",
    "on": "ustida",
    "to": "ga",
    "from": "dan",
    "of": "ning",
    "the": "",
    "a": "",
    "an": "",
}

PHRASES = {
    "you know": "bilasizmi",
    "i mean": "aytmoqchimanki",
    "thank you": "rahmat",
    "thanks": "rahmat",
    "good morning": "xayrli tong",
    "good night": "xayrli tun",
    "how are you": "qalaysiz",
    "what are you doing": "nima qilyapsiz",
    "let's go": "ketdik",
}

WORD_RE = re.compile(r"[^\W\d_]+(?:['`][^\W\d_]+)*", re.UNICODE)

# Maps an English part-of-speech (or Uzbek helper category) to an Uzbek label.
POS_UZ = {
    "pronoun": "Olmosh", "noun": "Ot", "verb": "Fe'l", "adjective": "Sifat",
    "adverb": "Ravish", "article": "Artikl", "preposition": "Predlog",
    "auxiliary": "Yordamchi fe'l", "particle": "Yuklama", "conjunction": "Bog'lovchi",
    "numeral": "Son", "interjection": "Undov",
    # Uzbek helper categories used as a fallback when no pos is returned.
    "komakchi": "Ko'makchi", "boglovchi": "Bog'lovchi",
}
POS_ORDER = [
    "Olmosh", "Ot", "Fe'l", "Sifat", "Ravish", "Son", "Artikl",
    "Predlog", "Ko'makchi", "Yordamchi fe'l", "Yuklama", "Bog'lovchi",
    "Undov", "Boshqa",
]


def pos_label(entry: dict[str, Any]) -> str:
    pos_raw = str(entry.get("pos", "")).lower()
    if not pos_raw and entry.get("helper"):
        pos_raw = str(entry["helper"]).lower()
    for key, label in POS_UZ.items():
        if key in pos_raw:
            return label
    return "Boshqa"


@dataclass
class Segment:
    start: float
    end: float
    text: str


@dataclass
class Word:
    start: float
    end: float
    word: str


def ensure_dirs() -> None:
    for path in (KINO_DIR, OUT_DIR, TMP_DIR, MODEL_DIR, DICT_DIR, CACHE_DIR):
        path.mkdir(parents=True, exist_ok=True)
    if load_dotenv and ENV_FILE.exists():
        load_dotenv(ENV_FILE)
    elif load_dotenv:
        load_dotenv()


def emit(kind: str, **data: Any) -> None:
    data["type"] = kind
    print(json.dumps(data, ensure_ascii=False), flush=True)


def _env_float(name: str, default: float) -> float:
    try:
        return float(os.getenv(name, str(default)))
    except (TypeError, ValueError):
        return default


def ai_throttle() -> None:
    """Small pause after a successful AI call to stay under provider rate limits.

    Paid providers (OpenAI/Gemini) tolerate a tiny gap; only Groq's free tier
    really needs it. Configurable via AI_THROTTLE_SEC (default 0.5s).
    """
    import time
    delay = _env_float("AI_THROTTLE_SEC", 0.5)
    if delay > 0:
        time.sleep(delay)


def ai_retry_backoff() -> None:
    """Pause before retrying after every provider failed (likely rate limited)."""
    import time
    delay = _env_float("AI_RETRY_BACKOFF_SEC", 6.0)
    if delay > 0:
        time.sleep(delay)


def run(cmd: list[str], cwd: Path | None = None) -> subprocess.CompletedProcess:
    return subprocess.run(
        cmd,
        cwd=str(cwd) if cwd else None,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding="utf-8",
        errors="replace",
    )


def require_tool(name: str) -> None:
    resolved = {"ffmpeg": FFMPEG, "ffprobe": FFPROBE}.get(name, name)
    if os.path.isfile(resolved) or shutil.which(resolved):
        return
    raise RuntimeError(f"{name} topilmadi. Uni dastur papkasiga yoki PATH ga qo'shing.")


def safe_stem(name: str) -> str:
    stem = Path(name).stem
    stem = re.sub(r"[^\w\-.]+", "_", stem, flags=re.UNICODE).strip("._")
    return stem or "video"


def is_url(value: str) -> bool:
    return bool(re.match(r"^https?://", (value or "").strip(), re.IGNORECASE))


def update_ytdlp() -> dict[str, Any]:
    """yt-dlp'ni o'zini-o'zi yangilaydi (`yt-dlp -U`). Kino/video saytlar tez-tez
    o'zgargani uchun yuklovchini yangi tutish yuklashning ishlab turishini ta'minlaydi.
    Frozen exe yonidagi tools/yt-dlp.exe binarisi yangilanadi."""
    if not (os.path.isfile(YTDLP) or shutil.which(YTDLP)):
        return {"updated": False, "message": "yt-dlp topilmadi"}
    try:
        proc = subprocess.run(
            [YTDLP, "-U"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=120,
        )
        out = ((proc.stdout or "") + (proc.stderr or "")).strip()
        low = out.lower()
        changed = "updating to" in low or "updated yt-dlp" in low
        return {"updated": changed, "message": out[-400:] if out else "yt-dlp allaqachon eng yangi"}
    except Exception as exc:  # noqa: BLE001 — yangilash majburiy emas, xatoni yumshoq qaytaramiz
        return {"updated": False, "message": f"Yangilab bo'lmadi: {exc}"}


def download_video(
    url: str,
    dest_dir: Path | None = None,
    use_cache: bool = True,
    progress_lo: float = 0.01,
    progress_hi: float = 0.05,
) -> Path:
    """yt-dlp qo'llab-quvvatlaydigan istalgan havoladan (YouTube, Instagram,
    kino saytlari va h.k.) videoni yuklab oladi.

    dest_dir=None bo'lsa — ishchi papkaga (qayta ishlash uchun, keshlanadi).
    Ko'rinadigan papka berilsa (masalan Videos/Yuklab olingan) — o'sha yerga toza
    nom bilan saqlaydi."""
    url = url.strip()
    if not (os.path.isfile(YTDLP) or shutil.which(YTDLP)):
        raise RuntimeError("yt-dlp topilmadi — havoladan yuklab bo'lmaydi.")

    dl_dir = dest_dir or (TMP_DIR / "yuklamalar")
    dl_dir.mkdir(parents=True, exist_ok=True)
    key = hashlib.sha1(url.encode("utf-8", "replace")).hexdigest()[:10]
    span = max(0.0, progress_hi - progress_lo)

    if use_cache:
        # Bir xil havola avval yuklangan bo'lsa — qaytadan yuklamaymiz.
        for existing in dl_dir.glob(f"*__{key}.*"):
            if existing.suffix.lower() in VIDEO_EXTS and existing.stat().st_size > 0:
                emit("progress", message="Video allaqachon yuklangan (keshdan)", progress=progress_lo)
                return existing
        out_tmpl = str(dl_dir / ("%(title).70s__" + key + ".%(ext)s"))
    else:
        out_tmpl = str(dl_dir / "%(title).120s.%(ext)s")

    emit("progress", message="Video havolasi tekshirilmoqda", progress=progress_lo)
    ffmpeg_dir = str(Path(FFMPEG).parent)
    cmd = [
        YTDLP, "--no-playlist", "--no-warnings", "--no-mtime",
        "-f", "bv*[height<=1080]+ba/b[height<=1080]/bv*+ba/b",
        "--merge-output-format", "mp4",
        "--ffmpeg-location", ffmpeg_dir,
        "--newline",
        "-o", out_tmpl,
        "--print", "after_move:filepath",
        url,
    ]
    # stderr'ni stdout'ga qo'shamiz — progress ba'zan stderr'da, ba'zan stdout'da
    # bo'ladi; bitta oqimda hammasini o'qib, ambiguity'dan qutulamiz.
    proc = subprocess.Popen(
        cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
        text=True, encoding="utf-8", errors="replace",
    )
    tail: list[str] = []
    final_path: Path | None = None
    last_pct = -1
    assert proc.stdout is not None
    for line in proc.stdout:
        line = line.rstrip("\r\n")
        tail.append(line)
        if len(tail) > 40:
            tail.pop(0)
        m = re.search(r"\[download\]\s+([\d.]+)%", line)
        if m:
            try:
                pct = int(float(m.group(1)))
            except ValueError:
                pct = last_pct
            if pct != last_pct:
                last_pct = pct
                emit("progress", message=f"Video yuklab olinmoqda {pct}%",
                     progress=progress_lo + span * (pct / 100.0))
        elif line.strip() and not line.lstrip().startswith("[") and (
            line.strip().lower().endswith((".mp4", ".mkv", ".webm", ".mov", ".m4v", ".avi"))
        ):
            final_path = Path(line.strip())
    proc.wait()

    if final_path is None or not final_path.exists():
        # Zaxira: papkadagi eng yangi mos faylni topamiz.
        pattern = f"*__{key}.*" if use_cache else "*"
        candidates = [p for p in dl_dir.glob(pattern) if p.suffix.lower() in VIDEO_EXTS and p.stat().st_size > 0]
        final_path = max(candidates, key=lambda p: p.stat().st_mtime) if candidates else None

    if proc.returncode != 0 or final_path is None or not final_path.exists():
        err = "\n".join(tail)[-600:] or "noma'lum"
        raise RuntimeError("Videoni yuklab bo'lmadi: " + err)
    emit("progress", message="Video yuklandi", progress=progress_hi)
    return final_path


def normalize_video_path(value: str) -> Path:
    if is_url(value):
        return download_video(value)
    path = Path(value)
    if not path.is_absolute():
        local = (ROOT / value).resolve()
        if local.exists():
            path = local
        else:
            path = KINO_DIR / value
    path = path.resolve()
    if not path.exists():
        raise RuntimeError(f"Video topilmadi: {path}")
    if path.suffix.lower() not in VIDEO_EXTS:
        raise RuntimeError("Bu video formati qo'llab-quvvatlanmaydi")
    return path


def scan(target_dir: str | None = None) -> list[dict[str, Any]]:
    ensure_dirs()
    scan_path = Path(target_dir) if target_dir else KINO_DIR
    if not scan_path.exists() or not scan_path.is_dir():
        return []

    videos: list[dict[str, Any]] = []
    try:
        for path in sorted(scan_path.iterdir()):
            if path.is_dir():
                videos.append({
                    "name": path.name,
                    "path": str(path),
                    "size": 0,
                    "subtitle": "",
                    "subtitleName": "",
                    "isDir": True,
                })
            elif path.is_file() and path.suffix.lower() in VIDEO_EXTS:
                sidecar = find_sidecar_subtitle(path)
                videos.append({
                    "name": path.name,
                    "path": str(path),
                    "size": path.stat().st_size,
                    "subtitle": str(sidecar) if sidecar else "",
                    "subtitleName": sidecar.name if sidecar else "",
                    "isDir": False,
                })
    except Exception:
        pass
    return videos


def find_sidecar_subtitle(video: Path) -> Path | None:
    for ext in SUB_EXTS:
        candidate = video.with_suffix(ext)
        if candidate.exists():
            return candidate
    for candidate in video.parent.glob(video.stem + ".*"):
        if candidate.suffix.lower() in SUB_EXTS:
            return candidate
    return None


def strip_tags(text: str) -> str:
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("{\\an8}", "").replace("{\\an2}", "")
    return " ".join(text.split())


def srt_time_to_seconds(raw: str) -> float:
    raw = raw.strip().replace(",", ".")
    hms = raw.split(":")
    if len(hms) == 3:
        h, m, s = hms
        return int(h) * 3600 + int(m) * 60 + float(s)
    if len(hms) == 2:
        m, s = hms
        return int(m) * 60 + float(s)
    return float(raw)


def seconds_to_srt_time(seconds: float) -> str:
    seconds = max(0.0, seconds)
    total_ms = int(round(seconds * 1000))
    ms = total_ms % 1000
    total_s = total_ms // 1000
    s = total_s % 60
    m = (total_s // 60) % 60
    h = total_s // 3600
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def seconds_to_ass_time(seconds: float) -> str:
    seconds = max(0.0, seconds)
    cs = int(round(seconds * 100))
    total_s = cs // 100
    cs %= 100
    s = total_s % 60
    m = (total_s // 60) % 60
    h = total_s // 3600
    return f"{h:d}:{m:02d}:{s:02d}.{cs:02d}"


def read_subtitle_text(path: Path) -> str:
    """Read a subtitle file, guessing the encoding.

    Sidecar SRTs for Russian films are frequently Windows-1251 (cp1251), which
    is invalid UTF-8. We try UTF-8 first, then common single-byte code pages,
    and finally latin-1 (which never fails) so text is never garbled.
    """
    data = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "cp1251", "cp1252", "cp1254"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("latin-1", errors="replace")


def parse_srt(path: Path) -> list[Segment]:
    raw = read_subtitle_text(path)
    raw = raw.replace("\r\n", "\n").replace("\r", "\n")
    blocks = re.split(r"\n\s*\n", raw.strip())
    segments: list[Segment] = []
    for block in blocks:
        lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue
        time_index = next((i for i, ln in enumerate(lines) if "-->" in ln), -1)
        if time_index < 0:
            continue
        start_raw, end_raw = [p.strip() for p in lines[time_index].split("-->", 1)]
        end_raw = end_raw.split()[0]
        text = strip_tags(" ".join(lines[time_index + 1:]))
        if not text:
            continue
        segments.append(Segment(srt_time_to_seconds(start_raw), srt_time_to_seconds(end_raw), text))
    return segments


def parse_vtt(path: Path) -> list[Segment]:
    raw = read_subtitle_text(path)
    raw = raw.replace("WEBVTT", "", 1)
    # Convert millisecond separators only inside timestamp lines so subtitle
    # text keeps its own periods (e.g. "Mr. Smith" stays intact).
    def _fix_timestamps(match: re.Match[str]) -> str:
        return match.group(0).replace(".", ",")

    raw = re.sub(r"\d{1,2}:\d{2}:\d{2}\.\d{3}|\d{2}:\d{2}\.\d{3}", _fix_timestamps, raw)
    tmp = TMP_DIR / (path.stem + "_from_vtt.srt")
    tmp.write_text(raw, encoding="utf-8")
    return parse_srt(tmp)


def write_srt(path: Path, segments: Iterable[Segment]) -> None:
    with path.open("w", encoding="utf-8-sig") as f:
        for i, seg in enumerate(segments, 1):
            f.write(f"{i}\n")
            f.write(f"{seconds_to_srt_time(seg.start)} --> {seconds_to_srt_time(seg.end)}\n")
            f.write(wrap_text(seg.text, 48) + "\n\n")


def probe_resolution(video: Path) -> tuple[int, int]:
    proc = run(
        [
            FFPROBE, "-v", "error",
            "-select_streams", "v:0",
            "-show_entries", "stream=width,height",
            "-of", "csv=p=0:s=x",
            str(video),
        ]
    )
    if proc.returncode == 0 and "x" in proc.stdout:
        try:
            w, h = proc.stdout.strip().split("x")[:2]
            return int(w), int(h)
        except Exception:
            pass
    return 1280, 720


def extract_embedded_subtitle(video: Path, tmp_dir: Path) -> Path | None:
    proc = run(
        [
            FFPROBE, "-v", "error",
            "-select_streams", "s",
            "-show_entries", "stream=index,codec_name",
            "-of", "json",
            str(video),
        ]
    )
    if proc.returncode != 0:
        return None
    try:
        data = json.loads(proc.stdout or "{}")
    except json.JSONDecodeError:
        return None
    streams = data.get("streams") or []
    if not streams:
        return None
    out = tmp_dir / "embedded.srt"
    proc = run([FFMPEG, "-y", "-i", str(video), "-map", "0:s:0", str(out)])
    if proc.returncode == 0 and out.exists() and out.stat().st_size > 0:
        return out
    return None


def probe_duration_seconds(path: Path) -> float:
    proc = run(
        [
            FFPROBE, "-v", "error",
            "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1",
            str(path),
        ]
    )
    try:
        return float((proc.stdout or "").strip())
    except (TypeError, ValueError):
        return 0.0


def extract_audio(video: Path, audio_path: Path) -> None:
    require_tool("ffmpeg")
    proc = run(
        [
            FFMPEG, "-y", "-i", str(video),
            "-vn", "-ac", "1", "-ar", "16000",
            "-c:a", "libopus", "-b:a", "12k",
            str(audio_path),
        ]
    )
    if proc.returncode != 0:
        raise RuntimeError("Audio ajratishda xato: " + (proc.stderr[-700:] or "noma'lum"))


def value(obj: Any, key: str, default: Any = None) -> Any:
    if isinstance(obj, dict):
        return obj.get(key, default)
    return getattr(obj, key, default)


def model_candidates(primary_env: str, default: str, fallback_env: str, fallbacks: list[str]) -> list[str]:
    # Bo'sh ("") env qiymati ham default sifatida qabul qilinadi — shunda ilova
    # meros bo'lgan noto'g'ri ANTHROPIC_MODEL kabi qiymatlarni "" bilan tozalab,
    # kod ichidagi standart modelga qaytishi mumkin bo'ladi.
    raw = [os.getenv(primary_env) or default, *os.getenv(fallback_env, "").split(","), *fallbacks]
    out: list[str] = []
    for item in raw:
        model = (item or "").strip()
        if model and model not in out:
            out.append(model)
    return out


def _groq_transcribe_file(
    client: Any,
    models: list[str],
    audio_path: Path,
    language: str,
    offset: float,
) -> tuple[list[Segment], list[Word], str]:
    """Transcribe a single audio file with Groq, shifting times by `offset`."""
    data = audio_path.read_bytes()
    resp = None
    last_error: Exception | None = None
    for model in models:
        kwargs: dict[str, Any] = {"model": model, "response_format": "verbose_json"}
        if language and language != "auto":
            kwargs["language"] = language
        try:
            try:
                resp = client.audio.transcriptions.create(
                    file=(audio_path.name, data),
                    timestamp_granularities=["word", "segment"],
                    **kwargs,
                )
            except TypeError:
                resp = client.audio.transcriptions.create(file=(audio_path.name, data), **kwargs)
            break
        except Exception as exc:
            last_error = exc
            continue
    if resp is None:
        raise RuntimeError(f"Groq Whisper ishlamadi: {last_error}")

    detected = str(value(resp, "language", "") or "").lower()
    segments: list[Segment] = []
    for seg in value(resp, "segments", []) or []:
        text = strip_tags(str(value(seg, "text", "") or ""))
        if text:
            segments.append(
                Segment(
                    offset + float(value(seg, "start", 0.0) or 0.0),
                    offset + float(value(seg, "end", 0.0) or 0.0),
                    text,
                )
            )
    words: list[Word] = []
    for item in value(resp, "words", []) or []:
        text = strip_tags(str(value(item, "word", "") or ""))
        if text:
            words.append(
                Word(
                    offset + float(value(item, "start", 0.0) or 0.0),
                    offset + float(value(item, "end", 0.0) or 0.0),
                    text,
                )
            )
    return segments, words, detected


def transcribe_with_groq(audio_path: Path, language: str) -> tuple[list[Segment], list[Word], str]:
    key = os.getenv("GROQ_API_KEY", "").strip()
    if not key:
        raise RuntimeError("GROQ_API_KEY topilmadi")
    from groq import Groq

    client = Groq(api_key=key)
    models = model_candidates(
        "GROQ_WHISPER_MODEL",
        "whisper-large-v3",
        "GROQ_WHISPER_FALLBACK_MODELS",
        ["whisper-large-v3-turbo"],
    )

    # Qisqa audio — bitta so'rovda. Uzun kino (Groq hajm limitiga yaqin yoki
    # undan katta) bo'laklarga bo'linadi va har bo'lak vaqti surib qo'shiladi.
    size = audio_path.stat().st_size
    size_limit = int(os.getenv("GROQ_AUDIO_CHUNK_BYTES", str(20 * 1024 * 1024)))
    if size <= size_limit:
        return _groq_transcribe_file(client, models, audio_path, language, 0.0)

    duration = probe_duration_seconds(audio_path)
    chunk_sec = max(60.0, float(os.getenv("GROQ_CHUNK_SECONDS", "1200")))  # 20 daqiqa
    if duration <= 0:
        duration = chunk_sec  # noma'lum — hech bo'lmasa bitta bo'lak

    total_chunks = int((duration + chunk_sec - 1) // chunk_sec)
    segments_all: list[Segment] = []
    words_all: list[Word] = []
    detected = ""
    idx = 0
    start = 0.0
    while start < duration - 0.05:
        chunk_path = audio_path.parent / f"chunk_{idx:03d}.ogg"
        # -ss so'rovdan oldin: tez qidirish, audio uchun aniqligi yetarli.
        proc = run(
            [
                FFMPEG, "-y", "-ss", f"{start:.3f}", "-t", f"{chunk_sec:.3f}",
                "-i", str(audio_path),
                "-vn", "-ac", "1", "-ar", "16000",
                "-c:a", "libopus", "-b:a", "12k",
                str(chunk_path),
            ]
        )
        if proc.returncode == 0 and chunk_path.exists() and chunk_path.stat().st_size > 0:
            emit(
                "progress",
                message=f"Groq transkripsiya (qism {idx + 1}/{total_chunks})",
                progress=0.30,
            )
            segs, wds, det = _groq_transcribe_file(client, models, chunk_path, language, start)
            segments_all.extend(segs)
            words_all.extend(wds)
            detected = detected or det
            chunk_path.unlink(missing_ok=True)
        idx += 1
        start += chunk_sec

    if not segments_all:
        raise RuntimeError("Groq Whisper bo'laklab transkripsiya qila olmadi")
    return segments_all, words_all, detected


def transcribe_with_faster_whisper(audio_path: Path, language: str) -> tuple[list[Segment], list[Word], str]:
    from faster_whisper import WhisperModel

    model_name = os.getenv("WHISPER_MODEL", "small")
    device = os.getenv("WHISPER_DEVICE", "auto")
    compute_type = os.getenv("WHISPER_COMPUTE_TYPE", "int8")
    model = WhisperModel(model_name, device=device, compute_type=compute_type, download_root=str(MODEL_DIR))
    kwargs: dict[str, Any] = {"word_timestamps": True}
    if language and language != "auto":
        kwargs["language"] = language
    seg_iter, info = model.transcribe(str(audio_path), **kwargs)
    segments: list[Segment] = []
    words: list[Word] = []
    for seg in seg_iter:
        text = strip_tags(seg.text or "")
        if text:
            segments.append(Segment(float(seg.start), float(seg.end), text))
        for w in seg.words or []:
            word = strip_tags(w.word or "")
            if word:
                words.append(Word(float(w.start), float(w.end), word))
    return segments, words, getattr(info, "language", "") or ""


def transcribe_with_whisper_cli(audio_path: Path, tmp_dir: Path, language: str) -> tuple[list[Segment], list[Word], str]:
    exe = shutil.which("whisper")
    if not exe:
        raise RuntimeError("whisper CLI topilmadi")
    model = os.getenv("WHISPER_MODEL", "small")
    cmd = [
        exe, str(audio_path),
        "--model", model,
        "--output_format", "srt",
        "--output_dir", str(tmp_dir),
    ]
    if language and language != "auto":
        cmd += ["--language", language]
    proc = run(cmd)
    if proc.returncode != 0:
        raise RuntimeError("whisper CLI xato: " + (proc.stderr[-700:] or "noma'lum"))
    srt = tmp_dir / (audio_path.stem + ".srt")
    segments = parse_srt(srt)
    return segments, words_from_segments(segments), language if language != "auto" else ""


def get_transcription(video: Path, tmp_dir: Path, source_lang: str) -> tuple[list[Segment], list[Word], str, str]:
    sidecar = find_sidecar_subtitle(video)
    if sidecar:
        emit("progress", message=f"SRT topildi: {sidecar.name}", progress=0.12)
        segments = parse_vtt(sidecar) if sidecar.suffix.lower() == ".vtt" else parse_srt(sidecar)
        return segments, words_from_segments(segments), source_lang if source_lang != "auto" else "", "sidecar_srt"

    embedded = extract_embedded_subtitle(video, tmp_dir)
    if embedded:
        emit("progress", message="Ichki subtitr ajratildi", progress=0.16)
        segments = parse_srt(embedded)
        return segments, words_from_segments(segments), source_lang if source_lang != "auto" else "", "embedded_srt"

    audio = tmp_dir / "audio.webm"
    emit("progress", message="Audio ajratilmoqda (yuqori siqilishda)", progress=0.18)
    extract_audio(video, audio)

    errors: list[str] = []
    if os.getenv("GROQ_API_KEY", "").strip():
        try:
            emit("progress", message="Groq Whisper transkripsiya qilmoqda", progress=0.30)
            segments, words, detected = transcribe_with_groq(audio, source_lang)
            return segments, words or words_from_segments(segments), detected, "groq"
        except Exception as exc:
            errors.append(f"Groq: {exc}")

    try:
        emit("progress", message="Lokal faster-whisper tekshirilmoqda", progress=0.30)
        segments, words, detected = transcribe_with_faster_whisper(audio, source_lang)
        return segments, words or words_from_segments(segments), detected, "faster_whisper"
    except Exception as exc:
        errors.append(f"faster-whisper: {exc}")

    try:
        emit("progress", message="Lokal whisper CLI tekshirilmoqda", progress=0.30)
        segments, words, detected = transcribe_with_whisper_cli(audio, tmp_dir, source_lang)
        return segments, words, detected, "whisper_cli"
    except Exception as exc:
        errors.append(f"whisper CLI: {exc}")

    detail = " | ".join(errors[-3:])
    raise RuntimeError(
        "Subtitr topilmadi va transkripsiya ishlamadi. Video yoniga .srt qo'ying "
        "yoki GROQ_API_KEY / faster-whisper / whisper CLI sozlang. " + detail
    )


# ---------------------------------------------------------------------------
# Checkpoint / cache — uzun kino uzilib qolsa noldan boshlamaslik uchun.
# Transkripsiya, tarjima va lug'at diskka saqlanadi va qayta ishlashda o'qiladi.
# ---------------------------------------------------------------------------

def cache_dir_for(video: Path) -> Path:
    try:
        st = video.stat()
        sig = f"{video.resolve()}|{st.st_size}|{int(st.st_mtime)}"
    except OSError:
        sig = str(video)
    key = hashlib.sha1(sig.encode("utf-8", "replace")).hexdigest()[:10]
    d = CACHE_DIR / f"{safe_stem(video.name)}_{key}"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _cache_enabled() -> bool:
    return os.getenv("SUBTITR_NO_CACHE", "") not in {"1", "true", "yes"}


def _load_json(path: Path) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None


def _save_json(path: Path, data: Any) -> None:
    try:
        path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    except OSError:
        pass


def load_transcription(cdir: Path) -> tuple[list[Segment], list[Word], str, str] | None:
    if not _cache_enabled():
        return None
    data = _load_json(cdir / "transcription.json")
    if not isinstance(data, dict) or not data.get("segments"):
        return None
    segs = [Segment(float(s[0]), float(s[1]), str(s[2])) for s in data["segments"]]
    words = [Word(float(w[0]), float(w[1]), str(w[2])) for w in data.get("words", [])]
    return segs, words, str(data.get("lang", "")), str(data.get("transcriber", "")) + "+cache"


def save_transcription(cdir: Path, segs: list[Segment], words: list[Word], lang: str, transcriber: str) -> None:
    if not _cache_enabled():
        return
    _save_json(cdir / "transcription.json", {
        "segments": [[s.start, s.end, s.text] for s in segs],
        "words": [[w.start, w.end, w.word] for w in words],
        "lang": lang,
        "transcriber": transcriber,
    })


def words_from_segments(segments: list[Segment]) -> list[Word]:
    out: list[Word] = []
    for seg in segments:
        items = WORD_RE.findall(seg.text)
        if not items:
            continue
        dur = max(0.1, seg.end - seg.start)
        slot = dur / len(items)
        for i, word in enumerate(items):
            start = seg.start + i * slot
            end = min(seg.end, start + max(0.18, slot * 0.8))
            out.append(Word(start, end, word))
    return out


def json_object_from_text(text: str) -> dict[str, Any]:
    text = text.strip()
    try:
        data = json.loads(text)
        return data if isinstance(data, dict) else {}
    except json.JSONDecodeError:
        pass
    match = re.search(r"\{.*\}", text, flags=re.S)
    if not match:
        return {}
    try:
        data = json.loads(match.group(0))
        return data if isinstance(data, dict) else {}
    except json.JSONDecodeError:
        return {}


def indexed_values(data: dict[str, Any], count: int, fallback: list[str]) -> tuple[list[str], list[str], list[str]]:
    block = data.get("translate")
    if not isinstance(block, dict):
        block = data
    out_tr: list[str] = []
    out_pos: list[str] = []
    out_lemma: list[str] = []
    for i in range(count):
        value = block.get(str(i)) if isinstance(block, dict) else None
        pos = ""
        lemma = ""
        if isinstance(value, dict):
            pos = str(value.get("pos") or "").lower().strip()
            lemma = str(value.get("lemma") or value.get("base") or "").strip()
            value = value.get("t") or value.get("translation") or value.get("text")
        text = str(value).strip() if value is not None else ""
        out_tr.append(text or fallback[i])
        out_pos.append(pos)
        out_lemma.append(lemma)
    return out_tr, out_pos, out_lemma


def subtitle_prompt(target_lang: str, source_lang: str, glossary: dict[str, str] | None = None) -> str:
    target_name = TARGET_LANG_NAMES.get(target_lang, target_lang)
    source_name = SOURCE_LANG_NAMES.get(source_lang, source_lang or "source language")
    gloss = ""
    if glossary:
        pairs = "; ".join(f"{k} = {v}" for k, v in list(glossary.items())[:60])
        gloss = (
            "Names/terms glossary — ALWAYS use exactly these target forms for "
            f"consistency across the whole film: {pairs}\n"
        )
    return (
        "You are an expert film subtitle translator.\n"
        f"Translate from {source_name} to natural spoken {target_name}.\n"
        + gloss +
        "Input is a JSON object with context_before, translate, and context_after.\n"
        "Use context_before/context_after only to understand pronouns, tone, and continuity.\n"
        "Return a flat JSON object for translate keys only, with exactly the same numeric keys.\n"
        "Rules:\n"
        "- Keep each subtitle concise, natural, and easy to read on screen.\n"
        "- Preserve names, brands, numbers, jokes, questions, warnings, and emotional tone.\n"
        "- Avoid word-for-word translation when it sounds unnatural.\n"
        "- Remove filler words only when they add no meaning.\n"
        "- For Uzbek, use clear everyday Uzbek Latin script with correct apostrophes.\n"
        "- No explanations, no markdown, JSON only."
    )


def ai_translate_batch(
    texts: list[str],
    target_lang: str,
    source_lang: str = "",
    before: list[str] | None = None,
    after: list[str] | None = None,
    glossary: dict[str, str] | None = None,
) -> tuple[list[str], str]:
    if not texts:
        return [], ""
    prompt = subtitle_prompt(target_lang, source_lang, glossary)
    payload = json.dumps(
        {
            "context_before": before or [],
            "translate": {str(i): text for i, text in enumerate(texts)},
            "context_after": after or [],
        },
        ensure_ascii=False,
    )

    providers = [
        ("openai", bool(os.getenv("OPENAI_API_KEY")), translate_openai),
        ("claude", bool(os.getenv("ANTHROPIC_API_KEY")), translate_claude),
        ("gemini", bool(os.getenv("GEMINI_API_KEY")), translate_gemini),
        ("groq", bool(os.getenv("GROQ_API_KEY")), translate_groq),
    ]
    errors: list[str] = []

    for attempt in range(5):
        for name, enabled, fn in providers:
            if not enabled:
                continue
            try:
                data = fn(prompt, payload)
                parsed = json_object_from_text(data)
                out_tr, _, _ = indexed_values(parsed, len(texts), texts)
                matched = sum(1 for i, value in enumerate(out_tr) if value and value != texts[i])
                if matched < max(1, len(texts) // 4):
                    raise RuntimeError("AI javobida tarjima kam")
                ai_throttle()
                return out_tr, name
            except Exception as exc:
                errors.append(f"{name}: {exc}")
                continue
        ai_retry_backoff()

    return [offline_translate_text(text) for text in texts], "offline_dictionary"


def translate_openai(prompt: str, payload: str) -> str:
    from openai import OpenAI

    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    last_error: Exception | None = None
    for model in model_candidates(
        "OPENAI_MODEL",
        "gpt-4o-mini",
        "OPENAI_FALLBACK_MODELS",
        ["gpt-4o"],
    ):
        try:
            if hasattr(client, "responses"):
                try:
                    resp = client.responses.create(
                        model=model,
                        input=[
                            {"role": "system", "content": prompt},
                            {"role": "user", "content": payload},
                        ],
                        text={"format": {"type": "json_object"}, "verbosity": "low"},
                        reasoning={"effort": os.getenv("OPENAI_REASONING_EFFORT", "low")},
                    )
                    text = getattr(resp, "output_text", "") or ""
                    if text:
                        return text
                except Exception as exc:
                    last_error = exc
            resp = client.chat.completions.create(
                model=model,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": payload},
                ],
            )
            return resp.choices[0].message.content or "{}"
        except Exception as exc:
            last_error = exc
            continue
    raise RuntimeError(f"OpenAI tarjima ishlamadi: {last_error}")


def translate_claude(prompt: str, payload: str) -> str:
    from anthropic import Anthropic

    client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    max_tokens = int(os.getenv("ANTHROPIC_MAX_TOKENS", "8192"))
    # Tarjima uchun "thinking" kerak emas — uni o'chirib token tejaymiz
    # (Sonnet 5 da u sukut bo'yicha YOQILGAN va qo'shimcha token yeydi).
    # ANTHROPIC_THINKING=adaptive qilib fikrlashni yoqish mumkin.
    thinking_off = os.getenv("ANTHROPIC_THINKING", "disabled").strip().lower() in {"disabled", "off", "none", ""}

    def _extract(resp: Any) -> str:
        return "".join(
            getattr(block, "text", "")
            for block in (resp.content or [])
            if getattr(block, "type", "") == "text"
        ).strip()

    last_error: Exception | None = None
    for model in model_candidates(
        "ANTHROPIC_MODEL",
        "claude-sonnet-5",
        "ANTHROPIC_FALLBACK_MODELS",
        ["claude-haiku-4-5"],
    ):
        base: dict[str, Any] = {
            "model": model,
            "max_tokens": max_tokens,
            "system": prompt,
            "messages": [{"role": "user", "content": payload}],
        }
        # Avval thinking'ni o'chirib urinamiz; ba'zi modellar (masalan Fable 5)
        # disabled'ni qabul qilmaydi — u holda thinking'siz zaxira urinish.
        attempts = [{**base, "thinking": {"type": "disabled"}}, base] if thinking_off else [base]
        for kwargs in attempts:
            try:
                text = _extract(client.messages.create(**kwargs))
                if text:
                    return text
            except Exception as exc:
                last_error = exc
                continue
    raise RuntimeError(f"Claude tarjima ishlamadi: {last_error}")


def translate_gemini(prompt: str, payload: str) -> str:
    from google import genai
    from google.genai import types

    client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
    last_error: Exception | None = None
    for model in model_candidates(
        "GEMINI_MODEL",
        "gemini-2.5-pro",
        "GEMINI_FALLBACK_MODELS",
        ["gemini-2.5-flash", "gemini-2.5-flash-lite"],
    ):
        try:
            resp = client.models.generate_content(
                model=model,
                contents=payload,
                config=types.GenerateContentConfig(
                    system_instruction=prompt,
                    temperature=0.15,
                    response_mime_type="application/json",
                ),
            )
            return resp.text or "{}"
        except Exception as exc:
            last_error = exc
            continue
    raise RuntimeError(f"Gemini tarjima ishlamadi: {last_error}")


def translate_groq(prompt: str, payload: str) -> str:
    from groq import Groq

    client = Groq(api_key=os.getenv("GROQ_API_KEY"))
    last_error: Exception | None = None
    for model in model_candidates(
        "GROQ_TEXT_MODEL",
        "llama-3.3-70b-versatile",
        "GROQ_TEXT_FALLBACK_MODELS",
        ["llama-3.1-8b-instant"],
    ):
        kwargs: dict[str, Any] = {
            "model": model,
            "messages": [
                {"role": "system", "content": prompt},
                {"role": "user", "content": payload},
            ],
            "temperature": 0.15,
        }
        try:
            try:
                resp = client.chat.completions.create(response_format={"type": "json_object"}, **kwargs)
            except Exception:
                resp = client.chat.completions.create(**kwargs)
            return resp.choices[0].message.content or "{}"
        except Exception as exc:
            last_error = exc
            continue
    raise RuntimeError(f"Groq tarjima ishlamadi: {last_error}")


def offline_translate_text(text: str) -> str:
    source = " ".join(text.split())
    if not source:
        return ""
    lower = source.lower()
    for phrase, translation in PHRASES.items():
        lower = re.sub(r"\b" + re.escape(phrase) + r"\b", translation, lower)
    parts: list[str] = []
    for token in re.findall(r"[A-Za-z']+|[^\w\s]+|\s+|\w+", source, re.UNICODE):
        if token.isspace() or re.fullmatch(r"[^\w\s]+", token):
            parts.append(token)
            continue
        key = token.lower().strip("'")
        translated = COMMON_WORDS.get(key)
        if translated is None:
            translated = token
        if token[:1].isupper() and translated:
            translated = translated[:1].upper() + translated[1:]
        parts.append(translated)
    result = " ".join("".join(parts).split())
    return result or source


def translate_segments(
    segments: list[Segment],
    target_lang: str,
    source_lang: str,
    progress_lo: float = 0.40,
    progress_hi: float = 0.56,
    cache_path: Path | None = None,
    glossary: dict[str, str] | None = None,
) -> tuple[list[Segment], str]:
    # Kattaroq batch = kamroq API chaqiruvi va kamroq takroriy prompt/kontekst
    # tokeni; sifat saqlanadi. Kontekst atigi bir necha satr — pronoun/oxang
    # uchun yetarli, lekin ortiqcha token yig'maydi. Env orqali sozlanadi.
    batch = max(10, int(os.getenv("TRANSLATE_BATCH", "50")))
    ctx = max(0, int(os.getenv("TRANSLATE_CONTEXT", "4")))
    total = len(segments)

    # Checkpoint: oldingi (uzilib qolgan) ishdan tayyor tarjimalarni yuklaymiz.
    cached: dict[str, str] = {}
    if cache_path is not None and _cache_enabled():
        data = _load_json(cache_path)
        if isinstance(data, dict):
            cached = {str(k): str(v) for k, v in data.items()}
    out_texts: list[str | None] = [cached.get(str(i)) for i in range(total)]
    provider = "cache" if any(t is not None for t in out_texts) else ""

    for start in range(0, total, batch):
        end_i = min(start + batch, total)
        idxs = list(range(start, end_i))
        done = end_i
        frac = done / total if total else 1.0
        if all(out_texts[i] is not None for i in idxs):
            emit("progress", message=f"Tarjima {done}/{total} (keshdan)",
                 progress=progress_lo + (progress_hi - progress_lo) * frac)
            continue
        chunk = segments[start:end_i]
        before = [seg.text for seg in segments[max(0, start - ctx):start]]
        after = [seg.text for seg in segments[end_i:end_i + ctx]]
        translated, provider = ai_translate_batch(
            [seg.text for seg in chunk],
            target_lang,
            source_lang=source_lang,
            before=before,
            after=after,
            glossary=glossary,
        )
        for j, i in enumerate(idxs):
            out_texts[i] = translated[j] if j < len(translated) else chunk[j].text
            cached[str(i)] = out_texts[i] or chunk[j].text
        if cache_path is not None and _cache_enabled():
            _save_json(cache_path, cached)  # har batchdan keyin saqlaymiz
        emit("progress", message=f"Tarjima {done}/{total}",
             progress=progress_lo + (progress_hi - progress_lo) * frac)

    out = [Segment(segments[i].start, segments[i].end, out_texts[i] or segments[i].text) for i in range(total)]
    return out, provider


def normalize_word(word: str) -> str:
    return word.lower().strip("`'\".,!?;:()[]{}")


def build_glossary(segments: list[Segment], target_lang: str, source_lang: str) -> dict[str, str]:
    """Film davomida izchillik uchun tez-tez uchraydigan ismlar/atamalarni
    bir marta tarjima qilib, {manba: nishon} lug'atini qaytaradi. AI kaliti
    bo'lmasa yoki ism topilmasa — bo'sh (izchillik faqat AI bilan mazmunli)."""
    has_ai = any(os.getenv(k) for k in ("OPENAI_API_KEY", "ANTHROPIC_API_KEY", "GEMINI_API_KEY", "GROQ_API_KEY"))
    if not has_ai:
        return {}
    # Bosh harfli (ism/atama nomzodi) tokenlarni chastota bo'yicha yig'amiz.
    freq: dict[str, int] = {}
    for seg in segments:
        for tok in WORD_RE.findall(seg.text):
            if len(tok) >= 2 and tok[:1].isupper() and not tok.isupper():
                freq[tok] = freq.get(tok, 0) + 1
    candidates = [w for w, c in sorted(freq.items(), key=lambda x: -x[1]) if c >= 3][:40]
    if not candidates:
        return {}
    try:
        translations, _pos, _lemma, _prov = translate_word_list(candidates, target_lang)
    except Exception:
        return {}
    glossary: dict[str, str] = {}
    for src, tgt in zip(candidates, translations):
        tgt = (tgt or "").strip()
        if tgt and tgt.lower() != src.lower():
            glossary[src] = tgt
    return glossary


# --- #2: Whisper "hallucination" (soxta matn) filtri ------------------------
# Whisper musiqa/jim sahnalarda takroriy yoki soxta matn chiqaradi. Ularni
# tozalaymiz (faqat transkripsiya natijasiga; tayyor .srt ga tegmaymiz).
HALLUCINATION_RE = [
    re.compile(r"amara\.org|opensubtitles|subscene|addic7ed", re.I),
    re.compile(r"subtitles?\s+by|субтитры\s+(?:подготовил|сделал|от)|редактор\s+субтитров", re.I),
    re.compile(r"продолжение\s+следует|thanks?\s+for\s+watching|thank\s+you\s+for\s+watching", re.I),
    re.compile(r"^\s*\[?\s*(music|музыка|музика|applause|аплодисменты|laughter|смех)\s*\]?\s*$", re.I),
    re.compile(r"^[\s♪♫🎵.,!?\-–—…]*$"),
]


def is_hallucination(text: str) -> bool:
    t = (text or "").strip()
    if len(t) < 1:
        return True
    return any(rx.search(t) for rx in HALLUCINATION_RE)


def clean_transcription(segments: list[Segment]) -> list[Segment]:
    """Soxta matn va ketma-ket takrorlangan qatorlarni olib tashlaydi."""
    out: list[Segment] = []
    prev_norm = None
    for seg in segments:
        text = " ".join((seg.text or "").split())
        if is_hallucination(text):
            continue
        norm = text.lower()
        if norm == prev_norm:  # ketma-ket bir xil (Whisper looping) — tashlab yuboramiz
            continue
        prev_norm = norm
        out.append(Segment(seg.start, seg.end, text))
    return out


# --- #3: subtitrlarni aqlli birlashtirish/bo'lish ---------------------------

def _split_text(text: str, limit: int) -> list[str]:
    """Uzun matnni gap/vergul/probel chegarasida <= limit bo'laklarga bo'ladi."""
    text = " ".join(text.split())
    if len(text) <= limit:
        return [text]
    parts: list[str] = []
    rest = text
    while len(rest) > limit:
        window = rest[:limit + 1]
        cut = -1
        for seps in (".!?", ",;:", " "):
            idxs = [window.rfind(s) for s in seps]
            cut = max(idxs)
            if cut > limit * 0.4:
                break
        if cut <= 0:
            cut = limit
        parts.append(rest[:cut + 1].strip())
        rest = rest[cut + 1:].strip()
    if rest:
        parts.append(rest)
    return [p for p in parts if p]


def polish_segments(segments: list[Segment]) -> list[Segment]:
    """Juda qisqa bo'laklarni birlashtiradi, juda uzunlarni tabiiy joyda bo'ladi."""
    if not segments:
        return segments
    max_chars = int(os.getenv("SUB_MAX_CHARS", "84"))
    min_dur = float(os.getenv("SUB_MERGE_MIN_DUR", "0.7"))
    min_chars = int(os.getenv("SUB_MERGE_MIN_CHARS", "12"))
    merge_gap = float(os.getenv("SUB_MERGE_GAP", "0.4"))

    # 1) Qisqa bo'laklarni keyingisi bilan birlashtirish.
    merged: list[Segment] = []
    i = 0
    n = len(segments)
    while i < n:
        seg = segments[i]
        while i + 1 < n:
            nxt = segments[i + 1]
            tiny = (seg.end - seg.start) < min_dur or len(seg.text) < min_chars
            close = (nxt.start - seg.end) < merge_gap
            fits = len(seg.text) + 1 + len(nxt.text) <= max_chars
            if tiny and close and fits:
                seg = Segment(seg.start, nxt.end, (seg.text.rstrip() + " " + nxt.text.lstrip()).strip())
                i += 1
            else:
                break
        merged.append(seg)
        i += 1

    # 2) Uzun bo'laklarni bo'lish (vaqtni belgilar soniga mutanosib taqsimlaymiz).
    out: list[Segment] = []
    for seg in merged:
        if len(seg.text) <= max_chars:
            out.append(seg)
            continue
        parts = _split_text(seg.text, max_chars)
        total = sum(len(p) for p in parts) or 1
        dur = seg.end - seg.start
        acc = 0
        for p in parts:
            f0 = acc / total
            acc += len(p)
            f1 = acc / total
            out.append(Segment(seg.start + dur * f0, seg.start + dur * f1, p))
    return out


def enforce_reading_speed(segments: list[Segment]) -> list[Segment]:
    """O'qishga qulay bo'lishi uchun subtitr davomiyligini moslashtiradi:
    minimal ko'rinish vaqti va belgi/sekund (CPS) chegarasi, keyingi subtitrga
    tegmagan holda oxirini cho'zadi. Vaqtlarni buzmaydi (faqat oxirini uzaytiradi)."""
    if not segments:
        return segments
    max_cps = float(os.getenv("SUB_MAX_CPS", "20"))
    min_dur = float(os.getenv("SUB_MIN_DUR", "1.0"))
    gap = 0.08
    out: list[Segment] = []
    for i, seg in enumerate(segments):
        start, end = seg.start, seg.end
        limit = segments[i + 1].start - gap if i + 1 < len(segments) else end + 6.0
        chars = len(seg.text)
        need_cps = chars / max_cps if max_cps > 0 else 0.0
        desired_end = start + max(min_dur, need_cps)
        end = min(max(end, desired_end), max(end, limit))
        if end < start:
            end = start
        out.append(Segment(start, end, seg.text))
    return out


def build_vocabulary(words: list[Word], source_lang: str, target_lang: str) -> list[dict[str, Any]]:
    # Chastota (#4) — har bir so'z necha marta uchraydi.
    freq: dict[str, int] = {}
    first: dict[str, Word] = {}
    for item in words:
        key = normalize_word(item.word)
        if len(key) < 2 or key.isdigit():
            continue
        freq[key] = freq.get(key, 0) + 1
        first.setdefault(key, item)
    keys = list(first.keys())
    translations, pos_list, lemma_list, provider = translate_word_list(keys, target_lang)
    entries: list[dict[str, Any]] = []
    for key, tr, pos, lemma in zip(keys, translations, pos_list, lemma_list):
        entries.append(
            {
                "word": key,
                "translation": tr,
                "pos": pos,
                "lemma": (lemma or "").strip().lower(),
                "count": freq.get(key, 1),
                "helper": helper_category(key, source_lang),
                "provider": provider,
            }
        )
    # Eng ko'p uchraganlar tepada.
    entries.sort(key=lambda e: -int(e.get("count", 1)))
    return entries


def dedupe_by_lemma(entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """So'zning turli shakllarini (kelgan/kelmoqda -> kelmoq) bitta yozuvга
    birlashtiradi va chastotalarni qo'shadi. Lemma bo'lmasa — so'zning o'zi."""
    groups: dict[str, dict[str, Any]] = {}
    order: list[str] = []
    for e in entries:
        lemma = (e.get("lemma") or e.get("word") or "").strip().lower() or str(e.get("word", ""))
        if lemma not in groups:
            groups[lemma] = {**e, "word": lemma, "count": int(e.get("count", 1)), "forms": {str(e.get("word", ""))}}
            order.append(lemma)
        else:
            g = groups[lemma]
            g["count"] = int(g.get("count", 1)) + int(e.get("count", 1))
            g["forms"].add(str(e.get("word", "")))
            if not g.get("translation") and e.get("translation"):
                g["translation"] = e["translation"]
    out = [groups[k] for k in order]
    out.sort(key=lambda e: -int(e.get("count", 1)))
    return out


def translate_word_list(words: list[str], target_lang: str) -> tuple[list[str], list[str], list[str], str]:
    if not words:
        return [], [], [], ""

    prompt = (
        "Translate each single word to "
        f"{TARGET_LANG_NAMES.get(target_lang, target_lang)}. Use the most common short meaning, "
        "1-3 words. For each word also give: 'pos' = part of speech in English "
        "(pronoun, noun, verb, adjective, adverb, article, preposition, conjunction, numeral, interjection); "
        "'lemma' = the dictionary/base form of the SOURCE word (e.g. plural/inflected -> base form). "
        "Return a JSON object where keys are the numeric indices and values are objects with 't', 'pos', 'lemma'. "
        "Example: {\"0\": {\"t\": \"olma\", \"pos\": \"noun\", \"lemma\": \"apple\"}}"
    )

    tr_list: list[str] = []
    pos_list: list[str] = []
    lemma_list: list[str] = []
    provider = ""

    batch_size = max(20, int(os.getenv("VOCAB_BATCH", "80")))
    for start in range(0, len(words), batch_size):
        chunk = words[start:start + batch_size]
        payload = json.dumps({str(i): word for i, word in enumerate(chunk)}, ensure_ascii=False)
        chunk_tr, chunk_pos, chunk_lemma, p = [], [], [], ""

        for attempt in range(5):
            for name, enabled, fn in [
                ("openai", bool(os.getenv("OPENAI_API_KEY")), translate_openai),
                ("claude", bool(os.getenv("ANTHROPIC_API_KEY")), translate_claude),
                ("gemini", bool(os.getenv("GEMINI_API_KEY")), translate_gemini),
                ("groq", bool(os.getenv("GROQ_API_KEY")), translate_groq),
            ]:
                if not enabled:
                    continue
                try:
                    parsed = json_object_from_text(fn(prompt, payload))
                    chunk_tr, chunk_pos, chunk_lemma = indexed_values(parsed, len(chunk), [offline_translate_word(word) for word in chunk])

                    matched = sum(1 for i, tr in enumerate(chunk_tr) if tr and tr.lower() != chunk[i].lower() and tr != offline_translate_word(chunk[i]))
                    if matched < max(1, len(chunk) // 5):
                        raise RuntimeError("Lug'atda yetarli tarjima qilinmadi")

                    p = name
                    break
                except Exception:
                    continue
            if chunk_tr:
                break
            ai_retry_backoff()

        if not chunk_tr:
            chunk_tr = [offline_translate_word(word) for word in chunk]
            chunk_pos = ["" for _ in chunk]
            chunk_lemma = ["" for _ in chunk]
            p = "offline_dictionary"

        tr_list.extend(chunk_tr)
        pos_list.extend(chunk_pos)
        lemma_list.extend(chunk_lemma)
        provider = p or provider

        # Kichik pauza — API rate-limitlariga tushib qolmaslik uchun (AI_THROTTLE_SEC bilan sozlanadi)
        ai_throttle()

    return tr_list, pos_list, lemma_list, provider


def offline_translate_word(word: str) -> str:
    return COMMON_WORDS.get(normalize_word(word), word)


def helper_category(word: str, source_lang: str) -> str:
    lang = (source_lang or "").lower()
    groups = HELPERS.get(lang, {})
    key = normalize_word(word)
    if lang == "ru":
        key = translit_ru_basic(key)
    for label, items in groups.items():
        if key in items:
            return label
    return ""


def translit_ru_basic(text: str) -> str:
    table = str.maketrans(
        {
            "а": "a", "б": "b", "в": "v", "г": "g", "д": "d", "е": "e",
            "ё": "e", "ж": "zh", "з": "z", "и": "i", "й": "y", "к": "k",
            "л": "l", "м": "m", "н": "n", "о": "o", "п": "p", "р": "r",
            "с": "s", "т": "t", "у": "u", "ф": "f", "х": "h", "ц": "ts",
            "ч": "ch", "ш": "sh", "щ": "sh", "ы": "i", "э": "e", "ю": "yu",
            "я": "ya", "ь": "", "ъ": "",
        }
    )
    return text.translate(table)


def ass_escape(text: str) -> str:
    return (text or "").replace("{", "(").replace("}", ")").replace("\n", r"\N")


def ass_color(hex_color: str, alpha: str = "00") -> str:
    h = hex_color.lstrip("#")
    if len(h) != 6:
        h = "FFFFFF"
    r, g, b = h[0:2], h[2:4], h[4:6]
    return f"&H{alpha}{b}{g}{r}".upper()


def inline_color(hex_color: str) -> str:
    h = hex_color.lstrip("#")
    if len(h) != 6:
        h = "FFFFFF"
    return "{\\c&H" + (h[4:6] + h[2:4] + h[0:2]).upper() + "&}"


def wrap_lines(text: str, limit: int, max_lines: int = 2) -> list[str]:
    if isinstance(text, list):
        text = " ".join(str(x) for x in text)
    if not isinstance(text, str):
        text = str(text)
    words = " ".join((text or "").split()).split()
    if not words:
        return []
    lines: list[str] = []
    current = ""
    for word in words:
        if current and len(current) + 1 + len(word) > limit and len(lines) < max_lines - 1:
            lines.append(current)
            current = word
        else:
            current = f"{current} {word}".strip()
    if current:
        lines.append(current)
    return lines[:max_lines]


def wrap_text(text: str, limit: int) -> str:
    return "\n".join(wrap_lines(text, limit, max_lines=3))


def layout_for(width: int, height: int, dual: bool, font_scale: float = 1.0) -> dict[str, int]:
    ar = width / max(1, height)
    base = width if ar < 0.85 else height
    font_size = max(16, round(base * (0.035 if not dual else 0.031) * font_scale))
    margin_lr = max(24, round(width * 0.06))
    cpl = max(18, min(48, int((width - margin_lr * 2) / (font_size * 0.52))))
    return {
        "font": font_size,
        "vocab_font": max(16, round(font_size * 0.9)),
        "margin_lr": margin_lr,
        "margin_v": max(28, round(height * (0.07 if ar >= 1.0 else 0.15))),
        "cpl": cpl,
        "outline": max(2, round(font_size * 0.09)),
    }


def ass_header(width: int, height: int, layout: dict[str, int], alignment: int = 2) -> str:
    return (
        "[Script Info]\n"
        "ScriptType: v4.00+\n"
        f"PlayResX: {width}\n"
        f"PlayResY: {height}\n"
        "WrapStyle: 0\n"
        "ScaledBorderAndShadow: yes\n\n"
        "[V4+ Styles]\n"
        "Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, "
        "BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, "
        "BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding\n"
        f"Style: Bottom,Noto Sans,{layout['font']},{ass_color('#FFFFFF')},&H000000FF,"
        f"{ass_color('#000000')},&H00000000,1,0,0,0,100,100,0,0,1,{layout['outline']},"
        f"{max(1, layout['outline'] // 2)},{alignment},{layout['margin_lr']},{layout['margin_lr']},"
        f"{layout['margin_v']},1\n"
        f"Style: Vocab,Noto Sans,{layout['vocab_font']},{ass_color('#FFFFFF')},&H000000FF,"
        f"{ass_color('#08111F')},{ass_color('#08111F', '55')},1,0,0,0,100,100,0,0,3,"
        f"{max(1, layout['outline'] // 2)},1,7,24,24,24,1\n\n"
        "[Events]\n"
        "Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text\n"
    )


def write_ass(
    path: Path,
    original: list[Segment],
    translated: list[Segment] | None,
    vocab_words: list[Word] | None,
    vocab_map: dict[str, str],
    width: int,
    height: int,
    font_scale: float = 1.0,
    position: str = "bottom",
    trans_color: str = "#FFE680",
) -> None:
    layout = layout_for(width, height, dual=translated is not None, font_scale=font_scale)
    alignment = 8 if position == "top" else 2  # ASS: 8=tepa markaz, 2=past markaz
    with path.open("w", encoding="utf-8") as f:
        f.write(ass_header(width, height, layout, alignment=alignment))
        for i, seg in enumerate(original):
            start = seconds_to_ass_time(seg.start)
            end = seconds_to_ass_time(seg.end)
            orig_lines = wrap_lines(seg.text, layout["cpl"] + (5 if translated else 0), 2)
            lines: list[str] = []
            if orig_lines:
                lines.append(inline_color("#FFFFFF") + r"\N".join(ass_escape(x) for x in orig_lines))
            if translated and i < len(translated):
                trans_lines = wrap_lines(translated[i].text, layout["cpl"] + 5, 2)
                if trans_lines:
                    lines.append(inline_color(trans_color) + r"\N".join(ass_escape(x) for x in trans_lines))
            if lines:
                joined = r"\N".join(lines)
                f.write(f"Dialogue: 0,{start},{end},Bottom,,0,0,0,,{joined}\n")

        if vocab_words:
            x = max(24, round(width * 0.04))
            base_y = round(height * 0.85)
            line_height = round(layout["vocab_font"] * 1.5)
            
            # Uzluksiz tepaga qarab harakatlanish (Scroll) mantig'i:
            # Ekranning 75% qismini bosib o'tadi
            distance = round(height * 0.75)
            target_y = base_y - distance
            duration_sec = 4.0
            
            # Tezlik = Masofa / Vaqt (piksellar / sekund)
            speed = distance / duration_sec
            
            # Bir-birining ustiga chiqib ketmasligi uchun 
            # ikkita so'z orasidagi minimal vaqt oralig'i
            min_time_gap = line_height / speed
            
            last_start_time = -999.0
            
            for item in vocab_words:
                key = normalize_word(item.word)
                tr = vocab_map.get(key, "")
                if not tr:
                    continue
                
                # So'z paydo bo'ladigan vaqtni hisoblash (oldingi so'zdan yetarlicha uzoqda bo'lishi kerak)
                actual_start = max(item.start, last_start_time + min_time_gap)
                actual_end = actual_start + duration_sec
                last_start_time = actual_start
                
                start = seconds_to_ass_time(actual_start)
                end = seconds_to_ass_time(actual_end)
                word = ass_escape(key)
                translation = ass_escape(tr)
                
                duration_ms = int(duration_sec * 1000)
                exit_start = duration_ms - 400
                exit_end = duration_ms
                
                # Chiroyli va zamonaviy chiqish/yo'qolish animatsiyasi (Pop-in va Pop-out)
                # \blur3, \fscx50, \fscy50 -> Boshida xira va kichik
                # \t(0,300) -> Tiniqlashib, 100% razmerga kattalashadi (chiroyli sakrab chiqish)
                # \t(exit_start,exit_end) -> Oxirida yana xiralashib kichrayadi
                override = "{\\bord1\\shad1\\fad(300,400)\\move(%d,%d,%d,%d)\\blur3\\fscx50\\fscy50\\t(0,300,\\blur0\\fscx100\\fscy100)\\t(%d,%d,\\blur3\\fscx50\\fscy50)}" % (
                    x, base_y, x, target_y, exit_start, exit_end
                )
                
                body = (
                    f"{override}{inline_color('#FFFFFF')}{word} - "
                    f"{inline_color('#7DD3FC')}{translation}"
                )
                f.write(f"Dialogue: 1,{start},{end},Vocab,,0,0,0,,{body}\n")


_HW_ENCODER_CACHE: str | None = None


def _encoder_works(codec: str) -> bool:
    proc = run(
        [
            FFMPEG, "-hide_banner", "-loglevel", "error",
            "-f", "lavfi", "-i", "color=c=black:s=128x128:d=0.1",
            "-c:v", codec, "-f", "null", "-",
        ]
    )
    return proc.returncode == 0


def pick_video_encoder() -> tuple[str, list[str], str]:
    """Return (codec, extra_ffmpeg_args, name) for the subtitle burn.

    Prefers a working GPU encoder (NVIDIA/Intel/AMD) — much faster on long
    films — falling back to CPU libx264. Configurable via SUB_ENCODER
    (auto|nvenc|qsv|amf|x264).
    """
    q = os.getenv("SUB_CRF", "25")
    x264 = ("libx264", ["-preset", os.getenv("SUB_PRESET", "veryfast"), "-crf", q], "libx264 (CPU)")
    presets: dict[str, tuple[str, list[str], str]] = {
        "nvenc": ("h264_nvenc", ["-preset", "p5", "-rc", "vbr", "-cq", q, "-b:v", "0"], "NVIDIA NVENC"),
        "qsv": ("h264_qsv", ["-global_quality", q, "-preset", "faster"], "Intel QSV"),
        "amf": ("h264_amf", ["-rc", "cqp", "-qp_i", q, "-qp_p", q], "AMD AMF"),
        "x264": x264,
    }
    choice = os.getenv("SUB_ENCODER", "auto").strip().lower()
    if choice in presets:
        return presets[choice]

    global _HW_ENCODER_CACHE
    if _HW_ENCODER_CACHE is None:
        _HW_ENCODER_CACHE = "x264"
        for name, codec in [("nvenc", "h264_nvenc"), ("qsv", "h264_qsv"), ("amf", "h264_amf")]:
            try:
                if _encoder_works(codec):
                    _HW_ENCODER_CACHE = name
                    break
            except Exception:
                continue
    return presets.get(_HW_ENCODER_CACHE, x264)


def burn_subtitles(
    video: Path,
    ass_path: Path,
    out_path: Path,
    progress_lo: float = 0.76,
    progress_hi: float = 0.98,
    label: str = "",
) -> None:
    require_tool("ffmpeg")
    total = probe_duration_seconds(video)
    suffix = f": {label}" if label else ""

    def _run(codec: str, venc_args: list[str], enc_name: str) -> tuple[int, str]:
        emit("progress", message=f"Video render qilinmoqda ({enc_name}){suffix}", progress=progress_lo)
        cmd = [
            FFMPEG, "-y",
            "-i", str(video),
            "-vf", f"ass={ass_path.name}",
            "-c:v", codec, *venc_args,
            "-c:a", "aac",
            "-b:a", "128k",
            "-movflags", "+faststart",
            "-progress", "pipe:1", "-nostats",
            str(out_path),
        ]
        # Streamli — render foizini jonli ko'rsatish uchun `-progress` (stdout)
        # o'qiymiz; stderr'ni alohida oqimda bo'shatamiz (aks holda buferi to'lib
        # ffmpeg'ni bloklaydi).
        proc = subprocess.Popen(
            cmd, cwd=str(ass_path.parent),
            stdout=subprocess.PIPE, stderr=subprocess.PIPE,
            text=True, encoding="utf-8", errors="replace",
        )
        stderr_chunks: list[str] = []

        def _drain_stderr() -> None:
            if proc.stderr is not None:
                for ln in proc.stderr:
                    stderr_chunks.append(ln)

        stderr_thread = threading.Thread(target=_drain_stderr, daemon=True)
        stderr_thread.start()
        last_pct = -1
        assert proc.stdout is not None
        for line in proc.stdout:
            line = line.strip()
            if line.startswith("out_time_us=") and total > 0:
                try:
                    secs = int(line.split("=", 1)[1]) / 1_000_000.0
                except ValueError:
                    continue
                frac = max(0.0, min(1.0, secs / total))
                pct = round(frac * 100)
                if pct != last_pct:
                    last_pct = pct
                    emit(
                        "progress",
                        message=f"Video render ({enc_name}){suffix} {pct}%",
                        progress=progress_lo + (progress_hi - progress_lo) * frac,
                    )
        proc.wait()
        stderr_thread.join(timeout=5)
        return proc.returncode, "".join(stderr_chunks)

    codec, venc_args, enc_name = pick_video_encoder()
    code, err = _run(codec, venc_args, enc_name)
    if code != 0 and codec != "libx264":
        # Apparat kodlash uzildi (drayver/format) — libx264 (CPU) ga qaytamiz.
        global _HW_ENCODER_CACHE
        _HW_ENCODER_CACHE = "x264"
        code, err = _run("libx264", ["-preset", os.getenv("SUB_PRESET", "veryfast"), "-crf", os.getenv("SUB_CRF", "25")], "libx264 (CPU)")
    if code != 0:
        raise RuntimeError("Video render xato: " + (err[-900:] or "noma'lum"))


def write_txt_transcript(path: Path, original: list[Segment], translated: list[Segment] | None) -> None:
    lines: list[str] = ["MATN (1-format: Aralash)", "======================", ""]
    for i, seg in enumerate(original):
        lines.append(seg.text)
        if translated and i < len(translated):
            lines.append(f"    {translated[i].text}")
        lines.append("")
        
    if translated:
        lines += ["", "MATN (2-format: Alohida)", "========================", ""]
        lines += ["--- ORIGINAL ---", ""]
        for seg in original:
            lines.append(seg.text)
            
        lines += ["", "--- TARJIMA ---", ""]
        for seg in translated:
            lines.append(seg.text)
            
    path.write_text("\n".join(lines), encoding="utf-8-sig")


def _vocab_freq_note(e: dict[str, Any]) -> str:
    c = int(e.get("count", 1) or 1)
    return f"  (x{c})" if c > 1 else ""


def write_txt_vocab(path: Path, entries: list[dict[str, Any]]) -> None:
    # Bir xil o'zakli so'zlarni birlashtiramiz (kelgan/keldi -> kelmoq).
    merged = dedupe_by_lemma(entries)
    lines = ["LUG'AT (1-format: Chastota bo'yicha)", "=================================", ""]
    helpers = [e for e in merged if e.get("helper")]
    main = [e for e in merged if not e.get("helper")]
    lines.append(f"Asosiy so'zlar ({len(main)}) — eng ko'p uchraganlar tepada")
    lines.append("-" * 40)
    for e in main:
        lines.append(f"{e['word']} - {e['translation']}{_vocab_freq_note(e)}")
    if helpers:
        lines += ["", "Yordamchi so'zlar", "-" * 40]
        for e in helpers:
            lines.append(f"{e['word']} - {e['translation']} ({e['helper']}){_vocab_freq_note(e)}")

    # Format 2
    lines += ["", "", "LUG'AT (2-format: So'z turkumlariga ajratilgan)", "===============================================", ""]
    grouped: dict[str, list[dict[str, Any]]] = {}
    for e in merged:
        grouped.setdefault(pos_label(e), []).append(e)

    for cat in POS_ORDER:
        items = grouped.get(cat, [])
        if items:
            lines += [f"\n{cat} ({len(items)})", "-" * 40]
            for e in items:
                lines.append(f"{e['word']} - {e['translation']}{_vocab_freq_note(e)}")

    path.write_text("\n".join(lines), encoding="utf-8-sig")


def write_docx_transcript(path: Path, original: list[Segment], translated: list[Segment] | None) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("DOCX uchun python-docx kerak: pip install python-docx") from exc
    doc = Document()
    doc.add_heading("Matn", level=1)
    for i, seg in enumerate(original):
        p = doc.add_paragraph()
        p.add_run(f"[{seconds_to_srt_time(seg.start)}] ").bold = True
        p.add_run(seg.text)
        if translated and i < len(translated):
            p2 = doc.add_paragraph(translated[i].text)
            p2.paragraph_format.left_indent = 240000
    doc.save(path)


def write_docx_vocab(path: Path, entries: list[dict[str, Any]]) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("DOCX uchun python-docx kerak: pip install python-docx") from exc
    merged = dedupe_by_lemma(entries)
    doc = Document()
    doc.add_heading("Lug'at", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "So'z"
    hdr[1].text = "Tarjima"
    hdr[2].text = "Turi"
    hdr[3].text = "Necha marta"
    for e in merged:
        row = table.add_row().cells
        row[0].text = str(e.get("word", ""))
        row[1].text = str(e.get("translation", ""))
        row[2].text = pos_label(e)
        row[3].text = str(int(e.get("count", 1) or 1))
    doc.save(path)


def _prepare_data(
    video_value: str,
    mode: str,
    source_lang: str,
    target_lang: str,
) -> dict[str, Any]:
    """Transkripsiya + (kerak bo'lsa) tarjima + lug'at. Renderlashga kerak
    bo'lgan hamma narsani lug'at (dict) qilib qaytaradi. `process()` (bir-martalik
    oqim) va `prepare` (tahrirlash oqimi) ikkalasi ham shuni ishlatadi — render
    logikasi ikki joyda takrorlanmaydi."""
    ensure_dirs()
    require_tool("ffmpeg")
    require_tool("ffprobe")
    video = normalize_video_path(video_value)
    stem = safe_stem(video.name)
    job_tmp = Path(tempfile.mkdtemp(prefix=stem + "_", dir=str(TMP_DIR)))

    cdir = cache_dir_for(video)
    try:
        emit("progress", message="Video o'qilmoqda", progress=0.05)
        # Checkpoint: oldingi ishdan transkripsiyani qayta ishlatamiz (Whisper
        # sekin/qimmat — takrorlamaslik uchun).
        cached_tr = load_transcription(cdir)
        if cached_tr is not None:
            original, words, detected_lang, transcriber = cached_tr
            emit("progress", message="Matn keshdan olindi", progress=0.35)
        else:
            original, words, detected_lang, transcriber = get_transcription(video, job_tmp, source_lang)
            if not original:
                raise RuntimeError("Nutq/subtitr topilmadi")
            save_transcription(cdir, original, words, detected_lang, transcriber)
        if not original:
            raise RuntimeError("Nutq/subtitr topilmadi")

        # #2 + #3: soxta matnni tozalash va aqlli birlashtirish/bo'lish.
        # (Whisper natijasiga qo'llaymiz; sozlash: SUBTITR_NO_POLISH=1 o'chiradi.)
        if os.getenv("SUBTITR_NO_POLISH", "") not in {"1", "true", "yes"}:
            cleaned = clean_transcription(original)
            original = polish_segments(cleaned) or original

        effective_lang = (source_lang if source_lang != "auto" else detected_lang) or "en"
        emit("progress", message=f"Matn tayyor ({transcriber})", progress=0.35)

        needs_translation = mode in {"dual", "dual_vocab", "srt", "transcript", "all"}
        translated: list[Segment] | None = None
        provider = ""
        if needs_translation:
            emit("progress", message="Ismlar izchilligi (glossariy)", progress=0.38)
            glossary = build_glossary(original, target_lang, effective_lang)
            emit("progress", message="Tarjima qilinmoqda", progress=0.40)
            tr_cache = cdir / f"translation_{effective_lang}_{target_lang}.json"
            translated, provider = translate_segments(
                original, target_lang, effective_lang,
                cache_path=tr_cache, glossary=glossary,
            )
            # #5 — o'qishga qulay vaqtlar (juda tez o'tib ketmasin).
            translated = enforce_reading_speed(translated)

        # #5 — original subtitr vaqtlarini ham o'qishga qulay qilamiz.
        display_original = enforce_reading_speed(original)

        VOCAB_MODES = {"vocabulary", "original_vocab", "dual_vocab", "all"}
        needs_vocab = mode in VOCAB_MODES
        entries: list[dict[str, Any]] = []
        if needs_vocab:
            emit("progress", message="Lug'at tuzilmoqda", progress=0.58)
            vocab_cache = cdir / f"vocab_{effective_lang}_{target_lang}.json"
            cached_vocab = _load_json(vocab_cache) if _cache_enabled() else None
            if isinstance(cached_vocab, list) and cached_vocab:
                entries = cached_vocab
            else:
                entries = build_vocabulary(words, effective_lang, target_lang)
                if _cache_enabled():
                    _save_json(vocab_cache, entries)

        width, height = probe_resolution(video)

        # Tahrirlash uchun original+tarjima bitta ro'yxatga birlashtiriladi
        # (write_ass tarjima qatorini indeks bo'yicha juftlaydi, o'z vaqtidan
        # foydalanmaydi — shuning uchun bu render uchun to'liq yetarli).
        segments: list[dict[str, Any]] = []
        for i, seg in enumerate(display_original):
            tr = translated[i].text if (translated and i < len(translated)) else ""
            segments.append({
                "start": round(seg.start, 3),
                "end": round(seg.end, 3),
                "original": seg.text,
                "translated": tr,
            })

        return {
            "video": str(video),
            "stem": stem,
            "mode": mode,
            "sourceLang": effective_lang,
            "targetLang": target_lang,
            "transcriber": transcriber,
            "translator": provider or "none",
            "width": width,
            "height": height,
            "segments": segments,
            "words": [[round(w.start, 3), round(w.end, 3), w.word] for w in words],
            "vocab": entries,
        }
    finally:
        if os.getenv("KEEP_DESKTOP_TMP", "") not in {"1", "true", "yes"}:
            shutil.rmtree(job_tmp, ignore_errors=True)


def _render_outputs(
    job: dict[str, Any],
    font_scale: float = 1.0,
    position: str = "bottom",
    sub_color: str = "#FFE680",
) -> dict[str, Any]:
    """`_prepare_data()` qaytargan (yoki foydalanuvchi tahrirlagan) `job`dan
    SRT/ASS/DOCX fayllar va subtitr kuydirilgan videoni tayyorlaydi."""
    ensure_dirs()
    require_tool("ffmpeg")
    require_tool("ffprobe")
    video = normalize_video_path(str(job["video"]))
    stem = safe_stem(str(job.get("stem") or video.name))
    mode = str(job["mode"])
    effective_lang = str(job.get("sourceLang") or "en")
    target_lang = str(job.get("targetLang") or "uz")
    transcriber = str(job.get("transcriber") or "")
    provider = str(job.get("translator") or "none")
    width = int(job.get("width") or 1280)
    height = int(job.get("height") or 720)

    seg_data = job.get("segments") or []
    display_original = [
        Segment(float(s["start"]), float(s["end"]), str(s.get("original", "")))
        for s in seg_data
    ]
    has_tr = any((str(s.get("translated") or "")).strip() for s in seg_data)
    translated: list[Segment] | None = (
        [Segment(float(s["start"]), float(s["end"]), str(s.get("translated") or "")) for s in seg_data]
        if has_tr else None
    )
    words = [Word(float(w[0]), float(w[1]), str(w[2])) for w in (job.get("words") or [])]
    entries: list[dict[str, Any]] = list(job.get("vocab") or [])
    vocab_map = {str(e["word"]): str(e["translation"]) for e in entries}

    emit("progress", message="Fayllar tayyorlanmoqda", progress=0.60)

    modes_to_render = [mode]
    if mode == "all":
        modes_to_render = ["dual_vocab", "original_vocab", "srt", "transcript", "vocabulary"]
    VOCAB_MODES = {"vocabulary", "original_vocab", "dual_vocab"}
    needs_vocab = any(m in VOCAB_MODES for m in modes_to_render)

    outputs: list[dict[str, str]] = []
    job_out_dir = OUT_DIR / stem
    job_out_dir.mkdir(parents=True, exist_ok=True)

    def add_output(kind: str, label: str, path: Path) -> None:
        outputs.append({"kind": kind, "label": label, "path": str(path), "name": path.name})

    if "srt" in modes_to_render:
        srt_orig = job_out_dir / f"{stem}_original.srt"
        write_srt(srt_orig, display_original)
        add_output("srt", "Original SRT", srt_orig)
        if translated:
            srt_uz = job_out_dir / f"{stem}_{target_lang}.srt"
            write_srt(srt_uz, translated)
            add_output("srt", "Tarjima SRT", srt_uz)

    if "transcript" in modes_to_render:
        txt = job_out_dir / f"{stem}_matn.txt"
        docx = job_out_dir / f"{stem}_matn.docx"
        write_txt_transcript(txt, display_original, translated)
        write_docx_transcript(docx, display_original, translated)
        add_output("txt", "Matn TXT", txt)
        add_output("docx", "Matn DOCX", docx)

    if "vocabulary" in modes_to_render:
        txt = job_out_dir / f"{stem}_lugat.txt"
        docx = job_out_dir / f"{stem}_lugat.docx"
        write_txt_vocab(txt, entries)
        write_docx_vocab(docx, entries)
        add_output("txt", "Lug'at TXT", txt)
        add_output("docx", "Lug'at DOCX", docx)

    render_modes = [m for m in modes_to_render if m in {"original", "dual", "original_vocab", "dual_vocab"}]
    n_render = max(1, len(render_modes))
    for i, render_mode in enumerate(render_modes):
        emit("progress", message=f"ASS tayyorlanmoqda: {render_mode}", progress=0.66)
        ass = job_out_dir / f"{stem}_{render_mode}.ass"
        out = job_out_dir / f"{stem}_{render_mode}.mp4"
        use_trans = translated if render_mode in {"dual", "dual_vocab"} else None
        use_words = words if render_mode in {"original_vocab", "dual_vocab"} else None
        write_ass(
            ass, display_original, use_trans, use_words, vocab_map, width, height,
            font_scale=font_scale, position=position, trans_color=sub_color,
        )
        add_output("ass", f"{render_mode} ASS", ass)
        # Render progressini bir necha video orasida bo'lib ko'rsatamiz.
        lo = 0.70 + (0.28 * i / n_render)
        hi = 0.70 + (0.28 * (i + 1) / n_render)
        burn_subtitles(video, ass, out, progress_lo=lo, progress_hi=hi, label=render_mode)
        add_output("video", f"{render_mode} video", out)

    # Vocab-li video rejimlar uchun lug'at faylini ham saqlab qo'yamiz.
    if needs_vocab and mode in {"dual_vocab", "original_vocab"}:
        txt = job_out_dir / f"{stem}_lugat.txt"
        if not txt.exists():
            write_txt_vocab(txt, entries)
        add_output("txt", "Lug'at TXT", txt)

    emit("progress", message="Yakunlandi", progress=1.0)
    return {
        "video": str(video),
        "mode": mode,
        "sourceLang": effective_lang,
        "targetLang": target_lang,
        "translator": provider,
        "transcriber": transcriber,
        "outputs": outputs,
        "outDir": str(job_out_dir),
    }


def _session_path_for(job: dict[str, Any]) -> Path:
    d = CACHE_DIR / "sessions"
    d.mkdir(parents=True, exist_ok=True)
    key = hashlib.sha1(str(job.get("video", "")).encode("utf-8", "replace")).hexdigest()[:10]
    return d / f"{safe_stem(str(job.get('stem') or 'session'))}_{key}.json"


def prepare_session(video_value: str, mode: str, source_lang: str, target_lang: str) -> dict[str, Any]:
    """Renderlashdan OLDIN transkripsiya+tarjimani tayyorlab, seans faylига saqlaydi.
    Foydalanuvchi tarjimalarni ko'rib/tahrirlab, keyin `render` bilan yakunlaydi."""
    job = _prepare_data(video_value, mode, source_lang, target_lang)
    session_path = _session_path_for(job)
    session_path.write_text(json.dumps(job, ensure_ascii=False), encoding="utf-8")
    return {
        "session": str(session_path),
        "video": job["video"],
        "mode": job["mode"],
        "sourceLang": job["sourceLang"],
        "targetLang": job["targetLang"],
        "transcriber": job["transcriber"],
        "translator": job["translator"],
        "segments": job["segments"],
        "vocabCount": len(job.get("vocab") or []),
    }


def render_session(
    session_path: str,
    segments_path: str | None = None,
    font_scale: float = 1.0,
    position: str = "bottom",
    sub_color: str = "#FFE680",
) -> dict[str, Any]:
    """Seans faylini (va agar berilgan bo'lsa, tahrirlangan segmentlarni) o'qib,
    videoni renderlaydi."""
    data = _load_json(Path(session_path))
    if not isinstance(data, dict) or not data.get("segments"):
        raise RuntimeError("Seans fayli topilmadi yoki buzilgan.")
    if segments_path:
        edited = _load_json(Path(segments_path))
        if isinstance(edited, list) and edited:
            # Foydalanuvchi tahrirlagan segmentlar butunlay almashtiradi.
            data["segments"] = edited
    return _render_outputs(data, font_scale=font_scale, position=position, sub_color=sub_color)


def process(
    video_value: str,
    mode: str,
    source_lang: str,
    target_lang: str,
    font_scale: float = 1.0,
    position: str = "bottom",
    sub_color: str = "#FFE680",
) -> dict[str, Any]:
    job = _prepare_data(video_value, mode, source_lang, target_lang)
    return _render_outputs(job, font_scale=font_scale, position=position, sub_color=sub_color)


def install_deps() -> None:
    if getattr(sys, 'frozen', False):
        emit("done", message="Dastur qadoqlangan, paketlar ichida mavjud")
        return
        
    packages = [
        "python-dotenv>=1.0.0",
        "groq>=0.11.0",
        "openai>=1.40.0",
        "google-genai>=1.0.0",
        "anthropic>=0.40.0",
    ]
    cmd = [sys.executable, "-m", "pip", "install", "--user", *packages]
    emit("progress", message="Python paketlar o'rnatilmoqda", progress=0.1)
    proc = run(cmd)
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr[-1000:] or "pip install xato")
    emit("done", message="Paketlar o'rnatildi")


def main() -> int:
    ensure_dirs()
    parser = argparse.ArgumentParser(description="Subtitr desktop local processor")
    sub = parser.add_subparsers(dest="command", required=True)

    p_scan = sub.add_parser("scan")
    p_scan.add_argument("--dir", default=None)
    sub.add_parser("install-deps")

    p_process = sub.add_parser("process")
    p_process.add_argument("--video", required=True)
    p_process.add_argument(
        "--mode",
        default="dual_vocab",
        choices=["dual_vocab", "original_vocab", "dual", "original", "srt", "transcript", "vocabulary", "all"],
    )
    p_process.add_argument("--source-lang", default="auto")
    p_process.add_argument("--target-lang", default="uz")
    p_process.add_argument("--font-scale", type=float, default=1.0)
    p_process.add_argument("--position", default="bottom", choices=["bottom", "top"])
    p_process.add_argument("--sub-color", default="#FFE680")

    p_download = sub.add_parser("download")
    p_download.add_argument("--url", required=True)

    sub.add_parser("update-ytdlp")

    # #1 — ikki bosqichli oqim: transkripsiya+tarjima (prepare) → tahrir → render.
    p_prepare = sub.add_parser("prepare")
    p_prepare.add_argument("--video", required=True)
    p_prepare.add_argument(
        "--mode",
        default="dual_vocab",
        choices=["dual_vocab", "original_vocab", "dual", "original", "srt", "transcript", "vocabulary", "all"],
    )
    p_prepare.add_argument("--source-lang", default="auto")
    p_prepare.add_argument("--target-lang", default="uz")

    p_render = sub.add_parser("render")
    p_render.add_argument("--session", required=True)
    p_render.add_argument("--segments", default=None)
    p_render.add_argument("--font-scale", type=float, default=1.0)
    p_render.add_argument("--position", default="bottom", choices=["bottom", "top"])
    p_render.add_argument("--sub-color", default="#FFE680")

    args = parser.parse_args()
    try:
        if args.command == "scan":
            emit("done", videos=scan(args.dir), kinoDir=str(KINO_DIR), outDir=str(OUT_DIR))
        elif args.command == "install-deps":
            install_deps()
        elif args.command == "download":
            dest = KINO_DIR / "Yuklab olingan"
            path = download_video(args.url, dest_dir=dest, use_cache=False, progress_lo=0.0, progress_hi=1.0)
            emit("done", path=str(path), name=path.name, dir=str(dest))
        elif args.command == "update-ytdlp":
            emit("done", **update_ytdlp())
        elif args.command == "prepare":
            result = prepare_session(args.video, args.mode, args.source_lang, args.target_lang)
            emit("done", **result)
        elif args.command == "render":
            result = render_session(
                args.session, segments_path=args.segments,
                font_scale=args.font_scale, position=args.position, sub_color=args.sub_color,
            )
            emit("done", **result)
        elif args.command == "process":
            result = process(
                args.video, args.mode, args.source_lang, args.target_lang,
                font_scale=args.font_scale, position=args.position, sub_color=args.sub_color,
            )
            emit("done", **result)
        return 0
    except Exception as exc:
        emit("error", message=str(exc))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
